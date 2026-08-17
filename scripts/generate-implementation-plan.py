#!/usr/bin/env python3
"""Generate the StackEDU web app implementation plan (DOCX)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "StackEDU-App-Implementation-Plan.docx"


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(10)
    doc.add_paragraph()
    return t


def phase(doc, number, title, subtitle, duration, blocks):
    """Render one implementation phase with a consistent shape."""
    h2(doc, f"Phase {number} — {title}")
    p = doc.add_paragraph()
    run = p.add_run(f"{subtitle}  ·  Estimated duration: {duration}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for label, items in blocks:
        h3(doc, label)
        bullets(doc, items)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_heading("StackEDU Web App — Implementation Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "AI-Powered Education Management System for Rwandan Tertiary Institutions\n"
        "Backend · Data Management · APIs · Security · Authentication · AI\n"
        "Planned deployment: app.stackedu.rw · app.stackedu.africa\n"
        "Frontend on Vercel · API on Render (Frankfurt) · Database on Neon (Frankfurt)\n"
        f"Prepared: {date.today().isoformat()} · StackForgeAI"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    para(
        doc,
        "This document is the living implementation plan. It explains what already "
        "exists, who owns each part of the system, and the order we build the remaining "
        "role portals. It is written in simple English so that both technical and "
        "non-technical readers can follow it.",
    )

    # ── 1. Executive summary ─────────────────────────────────────────────────
    h1(doc, "1. Executive Summary")

    para(
        doc,
        "StackEDU is a single system that runs the whole life of a student at a "
        "university or college: applying, being admitted, registering for courses, "
        "attending classes, paying fees, taking exams, getting results, using the "
        "library, and graduating. Today most Rwandan institutions do this with "
        "spreadsheets, paper, and disconnected tools. StackEDU replaces all of that "
        "with one login and one source of truth.",
    )

    h2(doc, "1.1 Where the project stands today")
    bullets(
        doc,
        [
            "Foundation is live: Hono API, Drizzle, one Postgres database per institution, "
            "session login, Cloudflare R2 uploads, and Resend email.",
            "Public admissions is live: apply, email verification before sign-in, documents, "
            "sandbox fee, submit, academic review, and tracking.",
            "The Student portal is live against real records: dashboard, courses, registration, "
            "results, transcript, fees, sandbox payment, receipt, library, notifications, "
            "timetable, onboarding, and assignment submit. StackEDU AI stays paused.",
            "The ICT Manager portal is live: users, access levels, revocation, audit log, "
            "institution settings, integrations with health checks, analytics, announcements, "
            "notifications, and account settings (notification preferences stored in the database).",
            "The Academic Admin portal is live: dashboard, applications inbox, students, courses "
            "and programmes (create and edit), calendar (create, edit, delete), timetable "
            "(read-only), faculty list, result approval, at-risk, reports, and notifications.",
            "Shared account layer is live on all portals: profile, password change (revokes all "
            "sessions), notification preferences in the database, TOTP two-factor authentication "
            "with authenticator apps, dark/light mode on this device, and a header notification "
            "dropdown with mark-as-read.",
            "Platform UX and security: institution name, logo, website, and location from ICT "
            "settings flow to login, apply, verification, emails, and receipts; ICT integration "
            "toggles gate Resend email and live MoMo/Airtel payments; three-minute inactivity "
            "prompt with one-minute logout countdown; app version v1.0.0 shown in the signed-in "
            "shell.",
            "Applicant email verification is live: registration creates the account but does not "
            "sign in. A six-digit code is emailed immediately; the applicant verifies in a modal "
            "on the create-account screen (or at /apply/verify when returning). Sign-in and access "
            "to the form, documents, and payment happen only after verification.",
            "Lecturer, Bursar, and Librarian screens still show mock data. They will write into "
            "the same tables the Student and Academic portals already read.",
            "Every signed-in page shows a short guide box so staff and students know what "
            "that screen is for.",
        ],
    )

    h2(doc, "1.4 Who owns what — and the build order")
    para(
        doc,
        "Roles are interconnected. We do not build each portal as its own fake system. "
        "Shared tables and APIs come first. Student, ICT, and Academic Admin are live. "
        "Bursar is next for fee holds that block registration. Academic work that students "
        "depend on — opening a semester, offerings, timetable — belongs to Academic Admin, "
        "not ICT.",
    )
    table(
        doc,
        ["Action a student needs", "Who must do it first", "Why"],
        [
            [
                "Sign in at all",
                "ICT Manager",
                "ICT creates and activates user accounts, roles, and access.",
            ],
            [
                "Register for courses",
                "Academic Admin, then Bursar if a hold exists",
                "Academic Admin opens the current semester and offerings. A fee hold from the Bursar blocks registration.",
            ],
            [
                "See a timetable / materials / assignments",
                "Academic Admin, then Lecturer",
                "Academic Admin assigns the lecturer and publishes the timetable. The lecturer uploads materials and assessments.",
            ],
            [
                "Pay fees and clear a hold",
                "Bursar",
                "Bursar owns fee structures, invoices, reconciliation, and fee holds.",
            ],
            [
                "See published results",
                "Lecturer, then Academic Admin",
                "Lecturer enters marks. Academic Admin publishes the batch.",
            ],
            [
                "Open e-library items",
                "Librarian",
                "Librarian publishes catalogue resources.",
            ],
            [
                "Activate a session / semester",
                "Academic Admin",
                "calendar.write — years, current semester, registration window. Not an ICT setting.",
            ],
        ],
    )
    para(
        doc,
        "Build order for the remaining portals: Bursar (fees and holds) → Lecturer "
        "(teaching and results entry) → Librarian (catalogue). ICT and Academic Admin "
        "are done. AI stays paused until the core portals are real.",
    )

    h2(doc, "1.2 The main recommendations in this document")
    table(
        doc,
        ["Question", "Recommendation", "Main reason"],
        [
            [
                "Which backend language?",
                "TypeScript on Node.js, using the Hono framework",
                "Same language as the web app and the future mobile app. One shared set of types. Already started in the repo.",
            ],
            [
                "Where does the backend run?",
                "Render, as a long-lived Node server in Frankfurt. The React frontend stays on "
                "Vercel.",
                "A permanent server holds one database connection pool, which survives the "
                "registration-day rush far better than serverless functions.",
            ],
            [
                "Which database?",
                "PostgreSQL on Neon (Frankfurt), one database per institution, through Drizzle ORM",
                "Academic data is highly relational. Neon sits in the same city as the API and "
                "autoscales with our uneven load. Separate databases mean one university can be "
                "restored without touching the others.",
            ],
            [
                "Which file storage?",
                "Cloudflare R2 for documents. Cloudflare Images for pictures.",
                "R2 charges nothing to download files. That matters a lot for an e-library.",
            ],
            [
                "Which login system?",
                "Better Auth, self-hosted inside our own database",
                "No cost per student, and personal data stays in our database, which helps with Rwandan law.",
            ],
            [
                "How do we add AI?",
                "Start with rules, then add machine learning. Use pgvector for library search.",
                "Rules give value on day one. Real machine learning needs real historical data first.",
            ],
        ],
    )

    h2(doc, "1.3 Headline timeline")
    para(
        doc,
        "With a team of four to five people, a working pilot for one institution is "
        "realistic in about five months. The full system, including AI and the "
        "Librarian portal, is realistic in about nine to eleven months. The mobile "
        "app follows after that.",
    )

    # ── 2. Codebase audit ────────────────────────────────────────────────────
    h1(doc, "2. What We Found in the Codebase")

    para(
        doc,
        "This section is a factual review of the app folder. It is important "
        "because the plan is built on top of what already exists, not on a "
        "blank page.",
    )

    h2(doc, "2.1 How the project is organised")
    para(
        doc,
        "The project is a 'monorepo'. That means several related projects live in one "
        "folder and share code. It is managed by Turborepo and the Bun package manager.",
    )
    table(
        doc,
        ["Folder", "What it holds", "Current state"],
        [
            ["apps/web", "The React web app that users see", "Almost complete visually"],
            ["apps/api", "The backend API", "Empty shell — one health endpoint"],
            ["packages/shared", "Types and validation rules shared by both", "Types written, validation empty"],
            ["design_system", "Colours, fonts, spacing, component samples", "Complete"],
            ["EMS-master-context.md", "The full product specification", "Complete and very useful"],
        ],
    )

    h2(doc, "2.2 The frontend technology already chosen")
    table(
        doc,
        ["Layer", "Technology", "Comment"],
        [
            ["Language", "TypeScript 5.8", "Strict typing across the whole project"],
            ["UI library", "React 19", "Function components and hooks only"],
            ["Routing", "TanStack Router v1", "One file per screen, auto-generated route tree"],
            ["Server data", "TanStack Query v5", "Ready for real API calls, currently unused"],
            ["Build tool", "Vite 6", "Fast builds"],
            ["Components", "shadcn/ui on Radix UI", "20 base components already copied in"],
            ["Styling", "Tailwind CSS v4", "Design tokens live in styles.css"],
            ["Forms", "React Hook Form + Zod", "Installed, ready to use"],
            ["Charts", "Recharts", "Used on dashboards"],
            ["Notifications", "Sonner toasts", "Already wired into the root route"],
        ],
    )

    h2(doc, "2.3 Screens that exist today")
    para(
        doc,
        "The table below counts the actual screen files found in the repository, and "
        "compares them against the screen list in the product specification.",
    )
    table(
        doc,
        ["Area", "Screens built", "Screens specified", "Status"],
        [
            ["Public authentication", "4", "4", "Complete"],
            ["Public admissions (apply)", "7", "Not itemised in spec", "Complete"],
            ["Student", "15", "15", "Complete (My Profile folded into other screens)"],
            ["Lecturer", "14", "13", "Complete"],
            ["Bursar", "9", "7", "Complete"],
            ["Academic Admin", "16", "13", "Complete"],
            ["ICT Manager", "15", "10", "Complete"],
            ["Librarian", "1 placeholder", "7", "Major gap — must be built"],
        ],
    )

    h2(doc, "2.4 The gaps we must close")
    numbered(
        doc,
        [
            "Lecturer, Bursar, and Librarian portals still use mock data files. Their screens "
            "exist but do not yet read or write the institution database.",
            "Academic timetable and full faculty management are read-only in the UI. Calendar, "
            "programmes, and courses can be created and edited; timetable conflict detection "
            "and lecturer assignment are not wired yet.",
            "Live payment gateways (MoMo, Airtel, DPO) and production webhooks are not connected. "
            "Student and admissions fees run in sandbox mode; ICT toggles block live MoMo/Airtel "
            "when turned off.",
            "Notification preferences are saved per user but not yet enforced when sending SMS or "
            "email beyond admissions.",
            "The central notification service (SMS via Pindo, templated email beyond admissions) "
            "is not built. In-app notification lists and the header dropdown are live for Student, "
            "ICT, and Academic.",
            "StackEDU AI features remain paused until core portals are complete.",
        ],
    )

    h2(doc, "2.5 The good news")
    bullets(
        doc,
        [
            "The data shapes are already agreed. packages/shared/src/types/ defines Student, "
            "Course, Result, Payment, LibraryResource, Application, and User. These become the "
            "database tables almost directly.",
            "Money is already handled correctly. Payment amounts are whole numbers of Rwandan "
            "Francs, never decimals. This avoids rounding bugs.",
            "Multi-tenancy is already anticipated. Every type carries an institutionId, which "
            "shows the original authors expected to serve many institutions. We go further and "
            "give each institution its own database, explained in section 4.2.",
            "The design system is finished, so no design work blocks development.",
            "The specification document is unusually complete, including permissions per role.",
        ],
    )

    # ── 3. Technology recommendations ────────────────────────────────────────
    h1(doc, "3. Technology Recommendations")

    h2(doc, "3.1 Backend programming language")

    para(
        doc,
        "This is the most important decision in the document, because it affects the "
        "web app, the future mobile app, hiring, and cost. Here is the honest "
        "comparison of the realistic options.",
    )

    table(
        doc,
        ["Option", "Strengths", "Weaknesses for StackEDU"],
        [
            [
                "TypeScript (Node.js)",
                "Same language as the web app and future React Native mobile app. Types and "
                "validation rules shared in one place. Huge developer pool in Kigali. Runs "
                "anywhere, including Render and Vercel. Already started in this repo.",
                "Not the fastest language for heavy number crunching. Manageable, because our "
                "heavy work is database queries, not computation.",
            ],
            [
                "Go",
                "Very fast and memory efficient. Excellent for high concurrency.",
                "Cannot share types with the frontend. Two languages to maintain. Smaller local "
                "hiring pool. Slower to build CRUD screens.",
            ],
            [
                "Python (FastAPI or Django)",
                "Best ecosystem for machine learning and data science.",
                "Cannot share types with the frontend. Two languages. Slower request handling on "
                "serverless. Better used as a small side service for AI only.",
            ],
            [
                "Java or Kotlin (Spring Boot)",
                "Very mature, common in large enterprises and banks.",
                "Heavy, slow to develop, expensive to host, poor fit for serverless. Overkill for "
                "this size of system.",
            ],
            [
                "C# (.NET)",
                "Strong tooling and performance.",
                "Same problems as Java: no type sharing, heavier hosting, smaller local pool.",
            ],
            [
                "PHP (Laravel)",
                "Fast to build, cheap hosting, large pool.",
                "No type sharing. Weaker fit for a modern API-first product feeding a mobile app.",
            ],
        ],
    )

    h3(doc, "Recommendation: TypeScript on Node.js, with the Hono framework")
    numbered(
        doc,
        [
            "One language everywhere. The web app, the backend, and the future mobile app all "
            "speak TypeScript. A developer can move between them in the same day.",
            "One definition of the truth. A Zod schema written once in packages/shared validates "
            "the form in the browser, validates the request in the API, and generates the API "
            "documentation. If a field changes, every part of the system fails to compile until "
            "it is fixed. That prevents a whole class of bugs.",
            "The mobile app becomes much cheaper. React Native and Expo reuse the same types, the "
            "same validation, and the same API client patterns.",
            "Hono is already installed in apps/api. There is nothing to throw away.",
            "Hono is portable. The same code runs on Render, Vercel, Cloudflare, and AWS Lambda "
            "without being rewritten. This is not theoretical: the hosting choice in this "
            "document changed after the API was already started, and it cost us nothing.",
            "Hiring in Kigali is realistic. JavaScript and TypeScript are the most widely known "
            "skills among local developers.",
        ],
    )

    para(
        doc,
        "One exception. If we later train a real machine learning model for at-risk "
        "student detection, we will add one small Python service just for that. It "
        "will run separately and talk to the main API over HTTP. This is planned in "
        "Phase 12 and does not change the main language choice.",
        bold=False,
    )

    h2(doc, "3.2 Hosting: frontend on Vercel, backend on Render")

    para(
        doc,
        "The specification assumed Cloudflare Workers. The decision now is a split. The "
        "React frontend stays on Vercel, where the marketing site already lives. The API "
        "runs on Render as a long-lived Node server in Frankfurt, with the database "
        "beside it on Neon Frankfurt.",
    )

    h3(doc, "Why split, instead of putting everything on Vercel")
    bullets(
        doc,
        [
            "The single biggest technical risk in this project is the registration window. Five "
            "thousand students try to register in the same hour, twice a year. Everything else "
            "is comfortable by comparison.",
            "On Vercel, that spike creates hundreds of separate function instances, and every "
            "one of them wants its own database connection. Handling this needs a connection "
            "pooler and careful tuning, and it is a well-known source of production incidents.",
            "On Render the API is one long-running program. It opens a pool of roughly 20 to 50 "
            "database connections once, then reuses them for every request. Under the same "
            "spike, its behaviour is far more predictable.",
            "Render includes background workers and scheduled jobs as a built-in feature, which "
            "simplifies Phase 13 and may remove the need for a separate job vendor at first.",
            "There is no 4.5 MB request limit and no function time limit, so long operations are "
            "less constrained.",
            "Pricing is a fixed monthly amount per service rather than usage-based, which is "
            "much easier to quote inside an institutional contract.",
        ],
    )

    h3(doc, "What we give up, and how we check it is acceptable")
    bullets(
        doc,
        [
            "Render has five regions: Oregon, Ohio, Virginia, Frankfurt, and Singapore. There is "
            "no African region, so Frankfurt is the closest option to Rwanda.",
            "Vercel does offer a Cape Town region, which is physically much closer to Kigali. "
            "That advantage is real but it is smaller than it appears.",
            "Internet traffic leaving Rwanda often travels north to Europe before reaching "
            "anywhere else, because of how the undersea cables and peering agreements work. "
            "Cape Town is therefore not automatically faster in practice.",
            "Required action in Phase 0: deploy a tiny test endpoint in both Frankfurt and Cape "
            "Town and measure the real round-trip time from a Kigali connection, on both MTN and "
            "Airtel mobile data. Decide with measurements, not assumptions.",
            "This must happen before anything is built, because a Render service's region cannot "
            "be changed later. Moving means creating a new service and migrating the data.",
            "If Frankfurt measures badly, the fallback is Vercel Cape Town with Supabase Cape "
            "Town. That option stays open until Phase 0 closes.",
        ],
    )

    h3(doc, "How the pieces connect")
    bullets(
        doc,
        [
            "The React app is served from Vercel's global network at app.stackedu.rw and "
            "app.stackedu.africa, so the screens themselves load quickly from a nearby location.",
            "The API runs on Render at api.stackedu.rw.",
            "Because the app and the API sit under the same registered domain (stackedu.rw), "
            "login cookies work across both. The cookie is issued with Domain=.stackedu.rw and "
            "SameSite=Lax.",
            "The API accepts cross-origin requests only from our own two app domains. Every "
            "other origin is refused.",
            "The API reaches Neon over an encrypted connection. Both sit in Frankfurt, so this "
            "hop costs only a few milliseconds.",
            "Unlike the marketing site, the app must NOT redirect users between .rw and .africa "
            "based on their country. A student who logs in on .rw must stay on .rw, or their "
            "session is lost.",
            "Each institution gets its own sub-path or subdomain later, for example "
            "app.stackedu.rw/ur. Data separation is handled by institutionId on every query "
            "regardless of the URL.",
        ],
    )

    h2(doc, "3.3 Database")

    para(
        doc,
        "Academic data is deeply relational. A result belongs to a student, a course, "
        "and a semester at the same time. A course has prerequisites which are other "
        "courses. A payment must never be lost or double counted. This rules out "
        "document databases such as MongoDB and points clearly to PostgreSQL.",
    )

    para(
        doc,
        "The rule that matters most here is colocation. A single page load fires between "
        "5 and 20 database queries, so database latency is multiplied on every screen. "
        "User-to-server latency is paid only once per request. Putting the API and the "
        "database in the same city is worth far more than putting either one closer to "
        "Rwanda. We must never split them across continents.",
    )

    table(
        doc,
        ["Provider", "Strengths", "Weaknesses"],
        [
            [
                "Neon (recommended)",
                "Compute scales up and down automatically, which suits our very uneven load. "
                "Database branching gives every pull request its own copy for safe migration "
                "testing. Built-in connection pooler, point-in-time restore, and read replicas "
                "for heavy reports. Standard PostgreSQL, so no lock-in.",
                "No African region — eight AWS regions, the nearest being Frankfurt. This is "
                "acceptable because the API is also in Frankfurt. Scale-to-zero must be turned "
                "off in production.",
            ],
            [
                "Supabase",
                "Has a Cape Town region, and bundles authentication, storage, and realtime.",
                "We deliberately chose Better Auth over its auth and Cloudflare R2 over its "
                "storage, so we would be using it for PostgreSQL alone. Its compute is a fixed "
                "instance size that you resize by hand, which fits our spiky load poorly. It "
                "would also be in the wrong city relative to a Render API.",
            ],
            [
                "Render PostgreSQL",
                "Same provider as the API, so it uses Render's private network with no public "
                "internet hop. Simplest possible setup.",
                "Fewer features than Neon: no branching, weaker autoscaling, and less tooling "
                "for a system of this size and lifespan.",
            ],
            [
                "AWS RDS",
                "Full control, and a Cape Town region is available.",
                "You manage backups, scaling, patching, and connection pooling yourself. More "
                "work and more cost for no benefit at our scale.",
            ],
        ],
    )

    h3(doc, "Recommendation")
    bullets(
        doc,
        [
            "PostgreSQL hosted on Neon, in the Frankfurt region (aws-eu-central-1) — the same "
            "city as the Render API.",
            "Accessed through Drizzle ORM, which is already installed. Drizzle writes plain SQL "
            "with full type safety and keeps migration files in version control.",
            "Autoscaling is the deciding feature. StackEDU is quiet for weeks, then five thousand "
            "students register in one hour. Neon's compute follows that curve by itself. A "
            "fixed-size instance means either paying for peak capacity all year, or being too "
            "small on the one day it matters.",
            "Database branching is the second reason. Across eighteen phases of schema changes, "
            "testing each migration against a real copy of the data is genuinely valuable.",
            "Switch OFF scale-to-zero in production. It saves money on idle databases but adds a "
            "delay to the first request after a quiet period, which users would feel.",
            "Use Neon's pooled connection string, and set the application's own pool to roughly "
            "20 to 50 connections.",
            "Enable the pgvector extension from day one, so semantic library search in Phase 12 "
            "needs no migration later.",
            "Add a read replica once reporting load grows, so a heavy bursar report can never "
            "slow down student registration.",
            "Each institution gets its own Neon database, with a small shared platform database "
            "alongside it. Section 4.2 explains why, and what it costs us.",
        ],
    )

    h3(doc, "Is this enough for 8 to 10 universities?")
    para(
        doc,
        "Comfortably, and it is worth showing the arithmetic. Ten institutions at around "
        "5,000 students each is 50,000 students. The largest tables grow at roughly 9 "
        "million attendance records a year, a few million notifications, and somewhere "
        "between 5 and 20 million audit log entries. Registrations, results, and payments "
        "are each well under a million a year. That totals roughly 20 to 35 million rows "
        "per year, so about 150 million after five years.",
    )
    para(
        doc,
        "For PostgreSQL that is a small-to-medium workload. A single well-indexed "
        "instance handles it without strain. The genuine risks are not total data volume "
        "at all — they are the concurrency spike at registration, running out of "
        "connections, and the audit and attendance tables eventually needing to be split "
        "by month. None of those are solved by picking a different provider; they are "
        "solved by the pooling, indexing, and partitioning work planned in Phase 1.",
    )

    h2(doc, "3.4 File and media management — the Cloudinary question")

    para(
        doc,
        "You asked whether Cloudinary should manage all files and media. The short "
        "answer is no — not because image optimisation is unimportant, but because "
        "documents and images are two different problems, and Cloudinary is priced for "
        "the wrong one. We use Cloudflare R2 for documents and Cloudflare Images for "
        "pictures. Here is the reasoning.",
    )

    h3(doc, "What files will the system actually hold?")
    table(
        doc,
        ["File type", "Typical size", "How often downloaded", "Example"],
        [
            ["Application documents", "1–10 MB", "A few times", "National ID, school certificate"],
            ["Passport photos", "under 2 MB", "Very often", "Student profile picture"],
            ["Course materials", "2–50 MB", "Often, by whole classes", "Lecture slides, PDFs"],
            ["E-library books and journals", "5–200 MB", "Very often", "E-books, research papers"],
            ["Assignment submissions", "1–20 MB", "Once or twice", "Student essays, code"],
            ["Generated documents", "under 1 MB", "Sometimes", "Transcripts, receipts"],
            ["Lecture video", "100 MB – 2 GB", "Often", "Recorded class sessions"],
            ["Institution branding", "under 1 MB", "On every page load", "Logos"],
        ],
    )

    h3(doc, "Why Cloudinary is the wrong choice for documents")
    numbered(
        doc,
        [
            "Cloudinary bills for bandwidth. An e-library is a download machine. If 3,000 "
            "students each download 200 MB of books in a semester, that is 600 GB of downloads. "
            "On a bandwidth-billed service this becomes the single largest bill in the project, "
            "and it grows every time you add an institution.",
            "Cloudinary is built for media transformation, not for document archives. Its best "
            "features (resize, crop, compress, format conversion) do nothing useful for a "
            "200-page PDF.",
            "Academic records must be kept for many years and must be auditable. That is a job "
            "for plain object storage with versioning and lifecycle rules.",
            "It is not S3-compatible in the standard way, so moving away later means rewriting "
            "the upload and download code.",
        ],
    )

    h3(doc, "Why not Cloudinary for the images either")
    para(
        doc,
        "Cloudinary would do the image job well. The reason we do not use it is that we "
        "are already on Cloudflare for R2, and Cloudflare Images does everything StackEDU "
        "actually needs from an image service: store the original, resize it per device, "
        "convert it to a modern format, and serve it from a global network. Cloudinary's "
        "genuine strengths — video pipelines, AI cropping, complex transformation chains — "
        "would go unused. Choosing it would mean a second vendor, a second bill, and a "
        "second set of credentials for no benefit we would feel.",
    )

    h3(doc, "Recommended approach: the right tool for each file type")
    table(
        doc,
        ["Service", "Used for", "Why"],
        [
            [
                "Cloudflare R2 (primary)",
                "All documents: application files, course materials, e-library books, assignment "
                "submissions, generated transcripts and receipts",
                "Charges nothing for downloads. Storage is cheap. It is S3-compatible, so any "
                "standard S3 library works and we are not locked in.",
            ],
            [
                "Cloudflare Images",
                "Pictures only: profile photos, passport photos, institution logos",
                "Automatically resizes and compresses per device and serves modern formats. A "
                "student on mobile data downloads a 30 KB photo instead of a 3 MB one, which "
                "genuinely matters in Rwanda. Same vendor and same account as R2.",
            ],
            [
                "Mux (optional, later)",
                "Recorded lecture streaming, if virtual classrooms are built",
                "Proper adaptive streaming so video quality drops instead of buffering on a weak "
                "connection. Only add this when video is actually launched.",
            ],
        ],
    )

    para(
        doc,
        "The result is one storage vendor instead of two, covering documents and images "
        "with a single account and a single bill, while keeping the zero-download-cost "
        "advantage that matters most for the e-library.",
    )

    para(
        doc,
        "Important note: using Cloudflare R2 and Cloudflare Images does not mean hosting "
        "the app on Cloudflare. They are storage and image delivery only. Our API signs "
        "upload links for them over the network. They work alongside Render and Vercel "
        "without any conflict.",
    )

    h2(doc, "3.5 Authentication")

    table(
        doc,
        ["Option", "Strengths", "Weaknesses"],
        [
            [
                "Better Auth (recommended)",
                "Free and self-hosted. All user data stays in our own PostgreSQL database. Has "
                "built-in two-factor authentication, one-time passwords, sessions, and password "
                "reset. Works with Drizzle. Issues tokens for the mobile app.",
                "We are responsible for configuring it correctly and keeping it updated.",
            ],
            [
                "Clerk",
                "Fastest to set up. Beautiful pre-built screens.",
                "Charges per active user each month. With 10,000 students across a few "
                "institutions this becomes very expensive. User data is stored outside Rwanda, "
                "which complicates compliance.",
            ],
            [
                "Supabase Auth",
                "Already included with the database. Good enough for common cases.",
                "Its role model is less flexible than what StackEDU needs, because we have six "
                "roles with strict module-level rules.",
            ],
        ],
    )

    para(
        doc,
        "Recommendation: Better Auth. The deciding factors are cost at student scale "
        "and data residency. A university with 8,000 students would pay a large monthly "
        "bill on a per-user service, and would also be sending every student's personal "
        "data to a third country.",
    )

    para(
        doc,
        "One honest caveat: Better Auth is a relatively young library. The risk is "
        "acceptable because our sessions and user records are ordinary rows in our own "
        "PostgreSQL database. If the library were ever abandoned, we would replace the "
        "library without losing any data. Pin the version and review updates "
        "deliberately.",
    )

    h2(doc, "3.6 Supporting services")

    table(
        doc,
        ["Need", "Recommendation", "Notes"],
        [
            ["Email", "Resend", "Simple API, good deliverability, works from serverless functions"],
            [
                "SMS",
                "Pindo (Rwanda) with Africa's Talking as backup",
                "Pindo is Kigali-based with direct local routes. Keep a second provider "
                "configured so SMS still works if one fails.",
            ],
            [
                "Mobile money",
                "MTN MoMo Collections API and Airtel Money API",
                "Direct integration means lower fees per transaction",
            ],
            [
                "Card payments",
                "DPO Pay",
                "Covers Rwanda and most of Africa. Flutterwave is a valid alternative that covers "
                "MoMo, Airtel, and cards in one integration — fewer integrations but higher fees.",
            ],
            [
                "Background jobs",
                "Graphile Worker — a job queue that lives inside our own PostgreSQL — run by a "
                "Render background worker",
                "The queue is a set of tables in the same database as the data. This allows "
                "transactional enqueue, explained below, and adds no vendor. Render's cron "
                "feature triggers the scheduled ones.",
            ],
            [
                "Caching and rate limits",
                "Render Key Value",
                "Redis-compatible, sits in the same Frankfurt region on Render's private "
                "network, so there is no public internet hop and no extra vendor.",
            ],
            ["Error tracking", "Sentry", "Catches crashes in both the web app and the API"],
            ["Uptime and logs", "Better Stack or Axiom", "Alerts if the app goes down at 2am"],
            [
                "PDF generation",
                "pdf-lib or @react-pdf/renderer",
                "Avoid headless Chrome. It is too slow and too large for serverless functions.",
            ],
            ["AI models", "Vercel AI SDK with OpenAI or Anthropic", "Called only from the server, never from the browser"],
        ],
    )

    h2(doc, "3.7 Mobile app strategy")

    para(
        doc,
        "The mobile app is not built in this plan, but every decision above is made so "
        "that it is cheap to add later.",
    )
    bullets(
        doc,
        [
            "Technology: React Native with Expo. The team already knows React.",
            "It will be added as apps/mobile inside the same monorepo.",
            "It imports packages/shared, so it automatically gets the same types and the same "
            "validation rules as the web app. No duplicated logic.",
            "This is why the API must support token-based login (Bearer tokens), not only "
            "browser cookies. Better Auth supports both.",
            "It is also why the API is versioned from day one. A student with an old app version "
            "installed must keep working after we deploy a change.",
            "The design tokens in design_system port directly to React Native styles.",
            "Expected saving: roughly 60 to 70 percent of the effort compared with building "
            "separate native apps for Android and iOS.",
        ],
    )

    h2(doc, "3.8 Final recommended stack")

    table(
        doc,
        ["Layer", "Choice"],
        [
            ["Frontend", "React 19, TypeScript, Vite, TanStack Router and Query, Tailwind v4, shadcn/ui"],
            ["Backend language", "TypeScript on Node.js 22"],
            ["API framework", "Hono"],
            ["Frontend hosting", "Vercel global edge — app.stackedu.rw and app.stackedu.africa"],
            ["Backend hosting", "Render web service, Frankfurt — api.stackedu.rw"],
            ["Database", "PostgreSQL on Neon (Frankfurt) with pgvector"],
            ["Database access", "Drizzle ORM"],
            ["Documents and files", "Cloudflare R2"],
            ["Images", "Cloudflare Images"],
            ["Authentication", "Better Auth, self-hosted"],
            ["Tenancy model", "One PostgreSQL database per institution, plus a small platform database"],
            ["Background jobs", "Graphile Worker inside PostgreSQL, run by a Render worker"],
            ["Cache and rate limiting", "Render Key Value"],
            ["Testing", "Vitest, Playwright, Testcontainers"],
            ["Email and SMS", "Resend, Pindo"],
            ["Payments", "MTN MoMo, Airtel Money, DPO Pay"],
            ["AI", "Vercel AI SDK, pgvector, optional Python service later"],
            ["Monitoring", "Sentry, Better Stack"],
            ["Mobile (future)", "React Native with Expo"],
        ],
    )

    # ── 4. Architecture ──────────────────────────────────────────────────────
    h1(doc, "4. System Architecture")

    h2(doc, "4.1 How a request travels through the system")
    numbered(
        doc,
        [
            "A student opens app.stackedu.rw on a phone. Vercel's global network serves the "
            "React app from the nearest location.",
            "The React app asks for data, for example GET "
            "https://api.stackedu.rw/v1/students/me/results, sending its session cookie.",
            "The request reaches the Render web service in Frankfurt. Render first checks the "
            "request comes from one of our two permitted app domains.",
            "Middleware checks the session token. If it is missing or expired, the request is "
            "rejected with 401.",
            "Middleware reads the user's role and institution from the session, then selects "
            "that institution's database connection pool.",
            "Role middleware checks whether this role is allowed to call this endpoint. If not, "
            "the request is rejected with 403.",
            "Zod validates the input. Bad input is rejected with 422 and a clear message.",
            "The route handler calls a service function. Services hold the business rules. "
            "Routes never talk to the database directly.",
            "The service queries Neon PostgreSQL through Drizzle, borrowing a connection from "
            "the server's shared pool rather than opening a new one. Every query is filtered by "
            "institutionId.",
            "If the action is significant, an entry is written to the audit log.",
            "If something slow must happen, such as sending 500 SMS messages, a job is pushed to "
            "the queue and the API replies immediately.",
            "The response is returned as JSON and cached by TanStack Query in the browser.",
        ],
    )

    h2(doc, "4.2 Keeping institutions separate: one database per institution")

    para(
        doc,
        "One StackEDU installation will serve several institutions. Their data must "
        "never mix. A bursar at one university must never see a student at another. "
        "There are two ways to achieve this, and the choice matters more than it first "
        "appears.",
    )

    table(
        doc,
        ["Approach", "How it works", "Verdict"],
        [
            [
                "Shared database",
                "All institutions share the same tables. Every row carries an institutionId, "
                "and every query filters on it.",
                "The usual answer for software with hundreds of small customers. Rejected here.",
            ],
            [
                "Database per institution (chosen)",
                "Each institution gets its own database. A small shared platform database holds "
                "only the list of institutions and the directory that routes a login to the "
                "right one.",
                "Chosen. With 8 to 10 large, high-value institutions, the isolation and restore "
                "benefits outweigh the extra operational work.",
            ],
        ],
    )

    h3(doc, "Why one database per institution")
    numbered(
        doc,
        [
            "Restores become surgical. This is the deciding argument. Imagine a registrar runs a "
            "bad bulk operation and wipes an intake's enrolments. With a shared database you "
            "must either roll everyone back to that moment, which destroys eight other "
            "institutions' work, or pick rows out of a backup by hand under pressure. With "
            "separate databases you restore that one institution in minutes and nobody else "
            "notices.",
            "Isolation stops being a promise and becomes a fact. A mistake in a query cannot "
            "reach another institution's data, because that data is not in the database the "
            "query is running against.",
            "One institution's busy day cannot slow another's. Registration week at a large "
            "university does not affect a smaller college on the platform.",
            "Ending a contract is clean. You hand over a database export rather than writing a "
            "script to extract one customer from a shared system.",
            "It wins procurement conversations. University committees ask whether their data is "
            "separated from other institutions. 'It is in its own database' is a far stronger "
            "answer than 'we filter by a column'.",
            "If one institution ever demands different data residency, only that institution has "
            "to move.",
        ],
    )

    h3(doc, "How it is structured")
    bullets(
        doc,
        [
            "One platform database, kept deliberately small. It holds the list of institutions, "
            "a directory mapping each email address to its institution, platform-level audit "
            "entries, and our own billing records. It contains no academic or financial data.",
            "One database per institution. It holds everything else: students, courses, results, "
            "payments, library resources, and that institution's own audit log.",
            "At login, the API looks up the email in the platform directory, learns which "
            "institution it belongs to, and from that point every query runs against that "
            "institution's database only.",
            "The institution is resolved from the session, never from the URL or the request "
            "body. This remains the single most important security rule in the system.",
            "Neon creates projects through its API, so adding a new institution is a script, not "
            "a manual setup task.",
            "Row Level Security stays enabled inside each database as a second layer of defence.",
        ],
    )

    h3(doc, "What this costs us, honestly")
    bullets(
        doc,
        [
            "Migrations must run across every database rather than one. This needs a small "
            "migration runner that loops over all institutions, reports progress, and stops "
            "safely if one fails. It is built once in Phase 1.",
            "Connection pooling needs more thought. The API keeps a pool per institution, opened "
            "when first needed and closed after a period of inactivity, rather than one large "
            "shared pool.",
            "Cross-institution reporting, for our own internal analytics, requires querying each "
            "database and combining the results. This is rare and can run as a nightly job.",
            "This approach suits 8 to 10 large institutions well. If StackEDU ever targets "
            "hundreds of small schools, the shared-database model should be reconsidered for "
            "that segment.",
        ],
    )

    h2(doc, "4.3 Environments")
    table(
        doc,
        ["Environment", "URL", "Database", "Purpose"],
        [
            [
                "Local",
                "localhost:5173 with the API on localhost",
                "Local PostgreSQL in Docker",
                "Day-to-day development",
            ],
            [
                "Preview",
                "Vercel preview URL plus a Render preview service, created per pull request",
                "A Neon branch created automatically for that pull request",
                "Review a change, and test its migration, before merging",
            ],
            [
                "Staging",
                "staging.app.stackedu.rw and staging-api.stackedu.rw",
                "A separate Neon organisation with its own platform and institution databases",
                "Client demos and user acceptance testing",
            ],
            [
                "Production",
                "app.stackedu.rw, app.stackedu.africa, api.stackedu.rw",
                "Platform database plus one database per institution, all in Neon Frankfurt",
                "Real users",
            ],
        ],
    )
    para(
        doc,
        "Real student data is never copied into staging. If realistic data is needed "
        "for a demo, it is generated or anonymised first. This is both a legal and a "
        "practical rule.",
    )

    # ── 5. Data management ───────────────────────────────────────────────────
    h1(doc, "5. Data Management Plan")

    h2(doc, "5.1 The database tables we need")
    para(
        doc,
        "These are grouped by area. The exact columns are finalised in Phase 1, but "
        "this is the complete list of what must exist.",
    )

    h3(doc, "The platform database (one, shared)")
    para(
        doc,
        "Deliberately small. It holds no academic or financial data, so a problem here "
        "cannot expose a student record.",
    )
    table(
        doc,
        ["Area", "Tables"],
        [
            [
                "Institution registry",
                "institutions, institution_databases (connection details), institution_status",
            ],
            [
                "Login routing",
                "user_directory (email to institution), login_attempts",
            ],
            [
                "Platform operations",
                "platform_audit_logs, platform_admins, billing_records, migration_history",
            ],
        ],
    )

    h3(doc, "Each institution's own database (one per institution)")
    para(
        doc,
        "Every institution gets an identical copy of this structure, holding only their "
        "own data.",
    )
    table(
        doc,
        ["Area", "Tables"],
        [
            [
                "Institution settings",
                "institution_settings, integrations, api_keys, audit_logs, "
                "notification_templates",
            ],
            [
                "People and access",
                "users, sessions, accounts, verification_tokens, roles, permissions, "
                "role_permissions, access_revocations",
            ],
            [
                "Students",
                "students, student_profiles, guardians, enrolments, enrolment_history, "
                "onboarding_progress",
            ],
            [
                "Admissions",
                "applications, application_documents, application_reviews, application_payments, "
                "admission_offers",
            ],
            [
                "Academic structure",
                "faculties, departments, programmes, programme_requirements, courses, "
                "course_prerequisites, academic_years, semesters, academic_calendar_events",
            ],
            [
                "Teaching and learning",
                "course_offerings, course_registrations, lecturer_assignments, timetable_slots, "
                "rooms, attendance_sessions, attendance_records, course_materials, announcements",
            ],
            [
                "Assessment",
                "assessments, assessment_components, submissions, submission_files, grades, "
                "results, result_batches, transcripts, gpa_snapshots",
            ],
            [
                "Finance",
                "fee_structures, fee_items, student_fee_accounts, invoices, payments, "
                "payment_attempts, receipts, fee_holds, reconciliation_records, refunds",
            ],
            [
                "Library",
                "library_resources, resource_files, collections, collection_items, "
                "resource_access_rules, resource_requests, resource_access_logs, "
                "resource_embeddings",
            ],
            [
                "Communication",
                "notifications, notification_deliveries, messages, message_threads",
            ],
            [
                "AI",
                "risk_scores, risk_factors, risk_interventions, ai_jobs, ai_usage_logs",
            ],
        ],
    )

    h2(doc, "5.2 Rules that apply to all data")
    bullets(
        doc,
        [
            "Every institution's database has exactly the same structure. There are no "
            "per-institution schema differences, because that would make migrations "
            "unmanageable. Differences in behaviour are handled by settings rows, not by "
            "different columns.",
            "Because each institution has its own database, most tables no longer need an "
            "institutionId column. It is kept only where a record must be traceable back to the "
            "platform level, such as billing.",
            "Money is always stored as a whole number of Rwandan Francs. Never a decimal. "
            "Decimals cause rounding errors that turn into disputes with students.",
            "Dates and times are stored in UTC with a timezone marker. They are displayed to "
            "users in Africa/Kigali time.",
            "Nothing important is ever truly deleted. Records get a deleted_at timestamp "
            "instead. Academic and financial history must be recoverable.",
            "Every table has created_at, updated_at, created_by, and updated_by.",
            "Human-facing identifiers follow a readable pattern: STU-2024-0481 for students, "
            "PAY-002841 for payments, APP-2024-1053 for applications. Internally the primary key "
            "is a UUID.",
            "Results, payments, and audit logs are append-only. A correction creates a new row "
            "that references the old one. The original is never overwritten.",
        ],
    )

    h2(doc, "5.3 Migrations")
    bullets(
        doc,
        [
            "All schema changes are made through Drizzle migration files that live in Git.",
            "Nobody changes the production database by hand. Ever.",
            "Because there is one database per institution, a migration runner applies each "
            "change across every institution in turn. It records progress in the platform "
            "database, so an interrupted run can be resumed rather than restarted.",
            "If a migration fails on one institution, the runner stops immediately. It never "
            "leaves half the platform on a new schema and half on the old one without telling "
            "anyone.",
            "The runner is built once in Phase 1 and then used for the rest of the project's "
            "life. It is worth investing in properly.",
            "Migrations run as a deliberate step during deployment, not silently on application "
            "start-up, so that a rollout can be paused if something looks wrong.",
            "Every migration must be tested against a copy of production data before release. "
            "Neon branching makes this straightforward.",
            "Destructive changes, such as dropping a column, are split into two releases: first "
            "stop using the column, then remove it in a later release.",
        ],
    )

    h2(doc, "5.4 Backups and recovery")
    table(
        doc,
        ["Item", "Policy"],
        [
            ["Automatic backups", "Daily, kept for 30 days, per institution"],
            [
                "Point-in-time recovery",
                "Any moment within the plan's retention window (typically 7 to 30 days — confirm "
                "at signup). Restores are performed per institution.",
            ],
            ["Off-site copy", "Weekly export of every database to Cloudflare R2 in a separate account"],
            ["File versioning", "Object versioning enabled on R2, so an overwritten file is recoverable"],
            ["Restore test", "Performed every quarter — a backup you have never restored is not a backup"],
            ["Recovery time objective", "4 hours to be back online"],
            [
                "Recovery point objective",
                "Near zero within the point-in-time window; up to one week if we ever fall back "
                "to the off-site export",
            ],
        ],
    )

    h3(doc, "Why one database per institution matters most here")
    para(
        doc,
        "This is where the tenancy decision in section 4.2 pays for itself. If one "
        "institution needs to be rolled back to yesterday morning, we restore only that "
        "institution. Every other institution keeps working, unaware anything happened. "
        "Under a shared database the same request would force a choice between rolling "
        "everyone back or extracting rows by hand during an incident, which is exactly "
        "the wrong time to be doing careful work.",
    )

    h3(doc, "Three failure scenarios and the response to each")
    table(
        doc,
        ["Scenario", "Response"],
        [
            [
                "Someone deletes or corrupts data by mistake",
                "Restore that institution to the moment before it, or branch the database at "
                "that timestamp to inspect the data first without touching production.",
            ],
            [
                "Neon has a problem in the Frankfurt region",
                "Rebuild from the weekly off-site export held in Cloudflare R2. This is the "
                "4-hour recovery target.",
            ],
            [
                "An account is compromised or a vendor is lost entirely",
                "The off-site export lives with a different vendor, under a separate account "
                "with separate credentials, precisely so it survives this.",
            ],
        ],
    )

    h2(doc, "5.5 Data retention and residency")
    bullets(
        doc,
        [
            "Academic records and transcripts: kept permanently. They are legal documents.",
            "Financial records: kept for at least 10 years for audit purposes.",
            "Audit logs: kept for 7 years.",
            "Rejected applications: kept for 2 years, then anonymised.",
            "Session and login records: kept for 90 days.",
            "Under Rwandan Law No. 058/2021, storing personal data outside Rwanda is only "
            "allowed if the data controller holds a valid registration certificate from the "
            "National Cyber Security Authority that authorises it. StackForgeAI must register "
            "before go-live. Registration is free and takes up to 30 working days.",
            "Because our server and database sit in Frankfurt, Germany, this registration is not "
            "optional, and the certificate must specifically permit storage outside Rwanda. "
            "Start it in Phase 0, not at the end.",
            "Hosting inside the European Union also brings GDPR into scope for our providers. "
            "Render and Neon both offer standard data processing agreements. Sign both and keep "
            "them on file, because institutions will ask to see them during procurement.",
        ],
    )

    h2(doc, "5.6 Migrating an institution's existing data")
    numbered(
        doc,
        [
            "Collect the institution's current spreadsheets and database exports.",
            "Map their columns to our tables and agree the mapping in writing with the registrar.",
            "Clean the data: remove duplicate students, fix invalid national ID numbers, "
            "standardise programme names.",
            "Provision that institution's own database, then import into a staging copy of it "
            "first and let the institution's own staff verify a sample.",
            "Run automatic checks: student counts match, fee totals match to the franc, no "
            "orphan records.",
            "Import into production during a planned quiet period.",
            "Keep the original files for 12 months in case anything must be traced back.",
        ],
    )

    # ── 6. API design ────────────────────────────────────────────────────────
    h1(doc, "6. API Design")

    h2(doc, "6.1 Ground rules")
    bullets(
        doc,
        [
            "REST over HTTPS, returning JSON. Simple, well understood, and easy for a mobile app "
            "to consume.",
            "Every path starts with /api/v1/. Version one is frozen once the mobile app ships. "
            "Breaking changes go into /api/v2/.",
            "Resource names are plural nouns: /students, /courses, /payments.",
            "Standard verbs only: GET to read, POST to create, PATCH to update part of a record, "
            "DELETE to archive.",
            "Every response has the same envelope, so the frontend handles all responses the "
            "same way.",
            "Every request carries a request ID that appears in the logs, so a support ticket "
            "can be traced to an exact line in the logs.",
        ],
    )

    h2(doc, "6.2 Standard responses")
    para(doc, "Success:")
    para(
        doc,
        '{ "data": { ... }, "meta": { "requestId": "req_abc123" } }',
    )
    para(doc, "List with paging:")
    para(
        doc,
        '{ "data": [ ... ], "meta": { "page": 1, "pageSize": 25, "total": 482, '
        '"totalPages": 20 } }',
    )
    para(doc, "Error:")
    para(
        doc,
        '{ "error": { "code": "FEE_HOLD_ACTIVE", "message": "You cannot register while '
        'a fee hold is on your account.", "details": { ... } } }',
    )
    para(
        doc,
        "Error messages are written for the person reading them, not for the developer. "
        "'FEE_HOLD_ACTIVE' is for the code; the message is for the student.",
    )

    h2(doc, "6.3 Status codes we use")
    table(
        doc,
        ["Code", "Meaning"],
        [
            ["200", "Success"],
            ["201", "Created"],
            ["204", "Success with nothing to return"],
            ["400", "The request itself is malformed"],
            ["401", "Not logged in, or the session expired"],
            ["403", "Logged in, but this role is not allowed to do this"],
            ["404", "Not found, or not visible to this institution"],
            ["409", "Conflict, for example registering twice for the same course"],
            ["422", "The data failed validation"],
            ["429", "Too many requests"],
            ["500", "Our fault — logged and alerted"],
        ],
    )

    h2(doc, "6.4 Endpoint map by area")

    table(
        doc,
        ["Area", "Main endpoints"],
        [
            [
                "Authentication",
                "POST /auth/login · POST /auth/verify-otp · POST /auth/logout · "
                "POST /auth/forgot-password · POST /auth/reset-password · GET /auth/session · "
                "POST /auth/refresh",
            ],
            [
                "Admissions (public)",
                "GET /apply/programmes · POST /apply/register · POST /apply/verify-email · "
                "POST /apply/resend-verification · GET/PATCH /apply/application · "
                "POST /apply/application/submit · document presign/confirm · "
                "POST /apply/payment",
            ],
            [
                "Admissions (staff)",
                "GET /admin/applications · GET /admin/applications/:id · "
                "POST /admin/applications/:id/decision · "
                "POST /admin/applications/:id/request-documents",
            ],
            [
                "Students",
                "GET /students · GET /students/:id · PATCH /students/:id · GET /students/me · "
                "GET /students/:id/timeline · POST /students/:id/status",
            ],
            [
                "Academic structure",
                "GET/POST /programmes · GET/POST /courses · GET/POST /semesters · "
                "GET/POST /academic-calendar · GET/POST /timetable",
            ],
            [
                "Registration",
                "GET /registration/window · GET /registration/available-courses · "
                "POST /registration · DELETE /registration/:id · POST /registration/submit",
            ],
            [
                "Teaching",
                "GET /lecturer/courses · GET /courses/:id/roster · "
                "POST /courses/:id/attendance · GET /courses/:id/attendance · "
                "POST /courses/:id/materials · POST /courses/:id/announcements",
            ],
            [
                "Assessment",
                "GET/POST /assessments · POST /submissions · GET /submissions/:id · "
                "POST /submissions/:id/grade · POST /results · POST /results/publish · "
                "GET /results/me · GET /transcripts/me",
            ],
            [
                "Finance",
                "GET/POST /fee-structures · GET /fees/me · GET /fees/students/:id · "
                "POST /payments/initiate · GET /payments/:id · GET /payments · "
                "POST /receipts/:id/void · POST /fee-holds · DELETE /fee-holds/:id · "
                "GET /reconciliation",
            ],
            [
                "Library",
                "GET /library/resources · POST /library/resources · "
                "PATCH /library/resources/:id · GET /library/resources/:id/download-url · "
                "GET/POST /library/collections · GET/POST /library/requests · "
                "GET /library/analytics · GET /library/search",
            ],
            [
                "Platform admin",
                "GET/POST /admin/users · PATCH /admin/users/:id · "
                "POST /admin/users/:id/revoke · GET /admin/audit-logs · "
                "GET/PATCH /admin/settings · GET/POST /admin/integrations · "
                "POST /admin/announcements",
            ],
            [
                "AI",
                "GET /ai/at-risk · POST /ai/at-risk/:id/intervention · "
                "POST /ai/timetable/generate · POST /ai/library/semantic-search",
            ],
            [
                "Files",
                "POST /files/upload-url · POST /files/confirm · GET /files/:id/download-url",
            ],
            [
                "Webhooks (incoming)",
                "POST /webhooks/momo · POST /webhooks/airtel · POST /webhooks/dpo · "
                "POST /webhooks/sms-status",
            ],
        ],
    )

    h2(doc, "6.5 Paging, filtering, and sorting")
    bullets(
        doc,
        [
            "Lists are always paged. Default page size 25, maximum 100. No endpoint ever returns "
            "every student at once.",
            "Filtering uses simple query parameters: ?programme=BSE&year=2&status=Active",
            "Sorting uses ?sort=-createdAt, where the minus sign means descending.",
            "Searching uses ?q=mugisha and runs against a PostgreSQL full-text index.",
            "For very long lists such as the audit log, cursor paging is used instead of page "
            "numbers, because it stays fast as the table grows.",
        ],
    )

    h2(doc, "6.6 Documentation")
    bullets(
        doc,
        [
            "The OpenAPI 3.1 specification is generated automatically from the Zod schemas using "
            "Hono's Zod OpenAPI integration. It cannot drift out of date, because it comes from "
            "the same code that validates the requests.",
            "Interactive documentation is published at /api/docs, protected by login on "
            "production.",
            "A typed API client is generated from the specification for the web app and later "
            "the mobile app.",
        ],
    )

    h2(doc, "6.7 Rate limiting")
    table(
        doc,
        ["Endpoint group", "Limit"],
        [
            ["Login and OTP", "5 attempts per 15 minutes per IP address and per account"],
            ["Password reset", "3 per hour per account"],
            ["Public application submission", "10 per hour per IP address"],
            ["Payment initiation", "10 per hour per student"],
            ["File upload links", "50 per hour per user"],
            ["General authenticated API", "1,000 per hour per user"],
        ],
    )

    # ── 7. Security ──────────────────────────────────────────────────────────
    h1(doc, "7. Security")

    h2(doc, "7.1 Role permission matrix")
    para(
        doc,
        "Y means full access, R means read only, and a dash means no access at all. "
        "This matrix is implemented as data in the database, not as scattered if-statements "
        "in the code, so it can be audited and changed without a deployment.",
    )
    table(
        doc,
        ["Module", "Student", "Lecturer", "Bursar", "Academic", "Librarian", "ICT"],
        [
            ["Own academic record", "R", "—", "—", "Y", "—", "Y"],
            ["All student records", "—", "R (own classes)", "—", "Y", "—", "Y"],
            ["Course registration", "Y (own)", "—", "—", "Y", "—", "Y"],
            ["Attendance", "R (own)", "Y (own courses)", "—", "Y", "—", "Y"],
            ["Result entry", "—", "Y (own courses)", "—", "Y", "—", "Y"],
            ["Result publishing", "—", "—", "—", "Y", "—", "Y"],
            ["Fees and payments", "R (own)", "—", "Y", "—", "—", "Y"],
            ["Fee holds", "—", "—", "Y", "R", "—", "Y"],
            ["Financial reports", "—", "—", "Y", "—", "—", "Y"],
            ["E-library content", "R", "R", "—", "—", "Y", "Y"],
            ["Course materials", "R (enrolled)", "Y (own courses)", "—", "R", "—", "Y"],
            ["User accounts", "—", "—", "—", "—", "—", "Y"],
            ["Audit log", "—", "—", "—", "—", "—", "Y"],
            ["System settings", "—", "—", "—", "—", "—", "Y"],
            ["At-risk alerts", "—", "Y (own courses)", "—", "Y", "—", "Y"],
        ],
    )

    para(
        doc,
        "Note the strict separation between admin levels. A bursar cannot see results. "
        "An academic admin cannot see payments. A librarian cannot see either. ICT's "
        "default job is the platform: users, roles, institution settings, integrations, "
        "and the audit log. ICT does not activate semesters, edit the course catalogue, "
        "or set fee structures — those stay with Academic Admin and the Bursar. Extra "
        "permissions can be granted on the Access Levels screen, but ownership does not "
        "move. The API must enforce this on every endpoint.",
    )

    h2(doc, "7.2 Protections we build in")
    table(
        doc,
        ["Risk", "Control"],
        [
            [
                "Someone reads another institution's data",
                "Each institution has its own database, so the data is not present to be read. "
                "The database is chosen from the session, never from the request. Row Level "
                "Security inside each database is a further layer.",
            ],
            [
                "A student changes their own grades",
                "Results are append-only. Only lecturers can enter, only academic admins can "
                "publish, every change is logged with who and when.",
            ],
            [
                "SQL injection",
                "Drizzle uses parameterised queries everywhere. Raw SQL strings are banned by "
                "lint rule.",
            ],
            [
                "Cross-site scripting",
                "React escapes output by default. A strict Content Security Policy header is "
                "set. Uploaded HTML is never rendered.",
            ],
            [
                "Stolen session token",
                "Short-lived access tokens (15 minutes) with refresh tokens. Tokens are revoked "
                "immediately when an ICT manager revokes access.",
            ],
            [
                "Brute force login",
                "Rate limits, account lockout after repeated failures, and mandatory two-factor "
                "authentication for all staff and admin accounts.",
            ],
            [
                "Malicious file upload",
                "File type and size checked before the signed link is issued, checked again "
                "after upload, virus scanned, and served from a separate domain so a malicious "
                "file cannot run as if it were part of our site.",
            ],
            [
                "Payment fraud or double charging",
                "Every payment carries an idempotency key. Webhook signatures are verified. "
                "Amounts are always re-checked on the server, never trusted from the browser.",
            ],
            [
                "Insider misuse",
                "Full audit log of every significant action, reviewed by the ICT manager. "
                "Access reviews every quarter.",
            ],
            [
                "Secrets leaking",
                "All keys stored as environment variables in Render and Vercel. Nothing "
                "sensitive in Git. Automatic secret scanning on every commit.",
            ],
            [
                "Denial of service",
                "Rate limiting inside the API, plus Cloudflare in front of api.stackedu.rw. "
                "Note that Vercel's firewall protects the frontend only, so the API needs its "
                "own protection.",
            ],
        ],
    )

    h2(doc, "7.3 Encryption")
    bullets(
        doc,
        [
            "In transit: HTTPS everywhere with TLS 1.3 and HSTS enabled.",
            "At rest: the database and object storage are encrypted by the provider.",
            "Extra field-level encryption for the most sensitive fields, such as national ID "
            "numbers, so that even a database dump does not expose them in plain text.",
            "Passwords are hashed with Argon2id. They are never stored or logged in readable "
            "form.",
        ],
    )

    h2(doc, "7.4 Compliance checklist for Rwanda")
    numbered(
        doc,
        [
            "Register StackForgeAI with the National Cyber Security Authority as a data "
            "controller and processor. Start this in Phase 0 because it takes up to 30 working "
            "days.",
            "Because data is stored in Frankfurt, the registration certificate must specifically "
            "authorise storage outside Rwanda, as required by Article 50 of Law No. 058/2021.",
            "Appoint a Data Protection Officer and publish their contact details.",
            "Keep a written record of what personal data we process and why.",
            "Sign a data processing agreement with each institution, making clear that the "
            "institution owns its data and StackForgeAI only processes it.",
            "Publish a privacy notice inside the app, in English, French, and Kinyarwanda.",
            "Build the ability for a person to request a copy of their data and to request "
            "correction.",
            "Have a written breach response plan, since breaches must be reported to the "
            "authority.",
            "Enable logging of data access, since failure to log is listed as an offence.",
        ],
    )

    h2(doc, "7.5 Security testing")
    bullets(
        doc,
        [
            "Automatic dependency scanning on every pull request.",
            "Static code analysis in the build pipeline.",
            "An internal penetration test before the pilot goes live.",
            "An external penetration test before the first full institution goes live.",
            "A written procedure for reporting vulnerabilities responsibly.",
        ],
    )

    # ── 8. Authentication detail ─────────────────────────────────────────────
    h1(doc, "8. Authentication and Access")

    h2(doc, "8.1 How a user logs in")
    numbered(
        doc,
        [
            "The user enters their email or student ID, plus their password.",
            "The server checks the password hash and confirms the account is active and belongs "
            "to a real institution.",
            "If two-factor authentication is switched on, a six-digit code is sent by SMS or "
            "email and the user is sent to the existing verify screen.",
            "The code is valid for 10 minutes, and only five attempts are allowed.",
            "On success the server issues an access token valid for 15 minutes and a refresh "
            "token valid for 7 days.",
            "On the web the tokens are stored in HTTP-only, secure, same-site cookies, which "
            "JavaScript cannot read. On mobile they are stored in the device's secure keychain.",
            "The user is sent to the dashboard for their role.",
        ],
    )

    h2(doc, "8.2 Rules per role")
    table(
        doc,
        ["Role", "Two-factor authentication", "Session length", "Notes"],
        [
            ["Student", "Optional but encouraged", "7 days", "Most users, mostly on mobile"],
            ["Lecturer", "Required", "24 hours", "Can change grades, so higher risk"],
            ["Bursar", "Required", "8 hours", "Handles money"],
            ["Academic Admin", "Required", "8 hours", "Publishes results"],
            ["Librarian", "Required", "24 hours", "Manages content"],
            ["ICT Manager", "Required, no exception", "4 hours", "Can do everything, so shortest session"],
        ],
    )

    h2(doc, "8.3 Password and account rules")
    bullets(
        doc,
        [
            "Minimum 10 characters. Checked against a list of commonly breached passwords.",
            "Password reset links are single-use and expire after 30 minutes.",
            "Accounts lock for 15 minutes after five failed attempts.",
            "New staff accounts are created by the ICT manager and the user sets their own "
            "password through an invitation link. Nobody sends passwords by email.",
            "Students receive their first login details by SMS at admission.",
            "When an ICT manager revokes access, every session for that user is destroyed within "
            "seconds, not at the next expiry.",
            "Revocation requires a written reason, which is stored in the audit log.",
        ],
    )

    h2(doc, "8.4 Replacing the stub in the code")
    para(
        doc,
        "Today the file routes/_auth.tsx contains 'const isAuthenticated = true'. In "
        "Phase 2 this becomes a real session check that runs before the screen loads, "
        "redirects to /login when there is no session, and also checks that the user's "
        "role matches the section they are trying to open. A student who types "
        "/bursar/ledger into the address bar must be sent away, not shown the page.",
    )

    # ── 9. File uploads ──────────────────────────────────────────────────────
    h1(doc, "9. File Upload and Media Pipeline")

    h2(doc, "9.1 Why files never pass through our API")
    para(
        doc,
        "Running the API on Render removes the 4.5 MB request limit that Vercel would "
        "have forced on us, so bypassing the API is now a deliberate choice rather than a "
        "workaround. We still make that choice, for four reasons.",
    )
    bullets(
        doc,
        [
            "Memory. A 200 MB e-book passing through the server occupies server memory while it "
            "travels. A few librarians uploading at the same time would exhaust it and slow down "
            "every other user on the system.",
            "Speed. The browser uploads straight to storage over a direct connection. Sending "
            "the file to Frankfurt first, only to forward it on, makes it slower for a user in "
            "Kigali.",
            "Cost. Every megabyte routed through the server is bandwidth we pay for twice, once "
            "arriving and once leaving.",
            "Reliability. A large upload that fails halfway leaves the server holding a partial "
            "file it must clean up. Storage services handle resumable uploads properly, which "
            "matters a great deal on an unstable mobile connection.",
        ],
    )

    h2(doc, "9.2 The upload flow, step by step")
    numbered(
        doc,
        [
            "The browser tells the API: 'I want to upload transcript.pdf, it is 8 MB, type "
            "application/pdf, for application APP-2024-1053.'",
            "The API checks that this user is allowed to upload for that record, that the file "
            "type is on the allowed list, and that the size is within the limit for that "
            "purpose.",
            "The API asks Cloudflare R2 for a pre-signed upload link. This link works for 15 "
            "minutes, for that exact file name and size, and nothing else.",
            "The API returns the link plus a file record ID with the status 'pending'.",
            "The browser uploads directly to R2 and shows a progress bar.",
            "The browser tells the API 'upload finished'.",
            "The API confirms with R2 that the file really exists and matches the expected size.",
            "A background job scans the file for viruses and, for PDFs, checks that it really is "
            "a PDF and not something renamed.",
            "The file record status becomes 'ready'. Only then does it appear in the interface.",
            "Any file that stays 'pending' for more than 24 hours is deleted automatically.",
        ],
    )

    h2(doc, "9.3 Download flow")
    bullets(
        doc,
        [
            "The browser asks the API for a download link.",
            "The API checks permission. For example, a library resource may be restricted to "
            "second-year engineering students only.",
            "The API returns a signed link valid for 5 minutes.",
            "The access is recorded, which feeds the library analytics screen.",
            "Files are served from a separate domain, for example files.stackedu.rw, so that a "
            "malicious upload can never run code in the context of the main app.",
        ],
    )

    h2(doc, "9.4 Limits per file purpose")
    table(
        doc,
        ["Purpose", "Allowed types", "Maximum size", "Stored in"],
        [
            ["Application documents", "PDF, JPG, PNG", "10 MB", "R2"],
            ["Passport photo", "JPG, PNG", "2 MB", "Cloudflare Images"],
            ["Profile picture", "JPG, PNG, WebP", "5 MB", "Cloudflare Images"],
            ["Course material", "PDF, DOCX, PPTX, XLSX, ZIP", "100 MB", "R2"],
            ["Assignment submission", "PDF, DOCX, ZIP, code archives", "50 MB", "R2"],
            ["E-library resource", "PDF, EPUB", "500 MB", "R2"],
            ["Institution logo", "PNG, SVG", "2 MB", "Cloudflare Images"],
            ["Lecture video (later)", "MP4, MOV", "5 GB", "Mux"],
        ],
    )

    h2(doc, "9.5 Keeping storage costs under control")
    bullets(
        doc,
        [
            "Files that have not been opened for 12 months move automatically to cheaper "
            "cold storage.",
            "Duplicate files are detected by checksum, so the same textbook uploaded twice is "
            "stored once.",
            "Cloudflare Images serves pictures in modern formats such as WebP and AVIF, at the exact "
            "size the screen needs. This typically cuts image data by 70 to 90 percent.",
            "Old assignment submissions are archived at the end of each academic year.",
            "A monthly storage report shows the largest consumers, so nothing grows unnoticed.",
        ],
    )

    # ── 10. AI ───────────────────────────────────────────────────────────────
    h1(doc, "10. AI Integration Plan")

    para(
        doc,
        "A word of honesty first. The specification calls the GPA calculator "
        "'AI-powered'. It is not artificial intelligence — it is arithmetic, and it "
        "must be exactly correct every time. We will build it as ordinary, thoroughly "
        "tested code. Calling deterministic maths 'AI' creates the wrong expectation "
        "and, worse, invites people to accept a wrong answer. The genuine AI features "
        "are listed below.",
    )

    h2(doc, "10.1 At-risk student detection")
    bullets(
        doc,
        [
            "Purpose: spot students who are heading towards failure early enough to help them.",
            "Stage one (Phase 12): a transparent rules engine. For example, attendance below 70 "
            "percent, or a drop of more than one grade point compared with last semester, or two "
            "missed assignments in a row, or no login for 21 days during term.",
            "Why rules first: they work from day one with no historical data, and a lecturer can "
            "see exactly why a student was flagged. Trust matters more than cleverness here.",
            "Stage two (after two semesters of real data): a gradient-boosted model trained on "
            "actual outcomes, which produces a risk score from 0 to 100.",
            "Every alert always shows its reasons. A black-box score that says 'this student "
            "will fail' with no explanation will be ignored by staff, and rightly so.",
            "Fairness check: the model must be tested to confirm it does not systematically flag "
            "students by gender, region, or economic background. If it does, it is corrected or "
            "not used.",
            "A human always decides the action. The system suggests; it never suspends anyone.",
        ],
    )

    h2(doc, "10.2 Semantic search in the e-library")
    bullets(
        doc,
        [
            "Problem: keyword search fails when a student searches for 'how buildings resist "
            "earthquakes' and the book is titled 'Seismic Structural Design'.",
            "Solution: convert every resource's title, abstract, and tags into a vector "
            "(a numerical fingerprint of meaning) and store it in PostgreSQL using pgvector.",
            "A student's search is converted the same way, and we find the closest matches by "
            "meaning rather than exact words.",
            "Results are then filtered by the student's access rights, so restricted resources "
            "never appear.",
            "This runs entirely in our own database. No book content is sent to a third party.",
        ],
    )

    h2(doc, "10.3 Timetable generation")
    bullets(
        doc,
        [
            "Problem: building a semester timetable by hand takes a registrar days and still "
            "produces clashes.",
            "Solution: a constraint solver, not a language model. The rules are hard facts: a "
            "lecturer cannot be in two rooms at once, a room has a fixed capacity, a student "
            "group cannot have two classes at the same time.",
            "The solver produces several valid options and the registrar chooses one.",
            "It also explains any constraint it could not satisfy, so a human can resolve it.",
        ],
    )

    h2(doc, "10.4 Document reading during admissions")
    bullets(
        doc,
        [
            "When an applicant uploads a school certificate, optical character recognition plus "
            "a language model extracts the name, subjects, and grades.",
            "The extracted values are pre-filled into the review screen, with the original "
            "document shown alongside.",
            "A staff member always confirms before anything is saved. The AI reduces typing; it "
            "does not make admission decisions.",
            "Expected saving: 5 to 10 minutes per application, which is significant during an "
            "intake of 2,000 applicants.",
        ],
    )

    h2(doc, "10.5 Student study assistant")
    bullets(
        doc,
        [
            "A chat assistant that answers questions strictly from the student's own course "
            "materials and the library resources they are allowed to see.",
            "It always cites the source document and page, so the student can verify it.",
            "It refuses to answer questions outside the course content, and it will not do "
            "assignments for students.",
            "Built as retrieval-augmented generation, using the same pgvector index as the "
            "library search.",
            "This is a Phase 12 stretch item, delivered only after the core system is stable.",
        ],
    )

    h2(doc, "10.6 Rules for all AI features")
    numbered(
        doc,
        [
            "AI models are called only from the server. API keys never reach the browser or the "
            "mobile app.",
            "Personal data is minimised before it is sent to any external model. Names and "
            "national ID numbers are stripped where the task does not need them.",
            "Every AI call is logged with its cost, so spending is visible and can be capped per "
            "institution.",
            "Every AI output that affects a person is reviewable and reversible by a human.",
            "If the AI provider is down, the feature degrades gracefully. The rest of the system "
            "keeps working. Nothing critical depends on an external model being available.",
            "Institutions can switch AI features off completely if they choose.",
        ],
    )

    # ── 11. Notifications ────────────────────────────────────────────────────
    h1(doc, "11. Notifications")

    h2(doc, "11.1 Channels")
    table(
        doc,
        ["Channel", "Used for", "Provider"],
        [
            ["In-app", "Everything. Always the record of truth.", "Our own database"],
            ["SMS", "Urgent and high-value events: results published, payment confirmed, fee deadline", "Pindo, Africa's Talking as backup"],
            ["Email", "Detailed items: admission letters, receipts, reports", "Resend"],
            ["Push (mobile, later)", "Same as SMS, but free", "Expo Push"],
        ],
    )

    para(
        doc,
        "SMS costs real money per message. The system must never send an SMS that could "
        "have been an in-app notification. Every SMS template is reviewed for necessity "
        "before it is switched on.",
    )

    h2(doc, "11.2 Events that trigger a notification")
    table(
        doc,
        ["Event", "Who is told", "Channels"],
        [
            ["Application received", "Applicant", "SMS + Email"],
            ["Documents requested", "Applicant", "SMS + Email"],
            ["Admission decision", "Applicant", "SMS + Email"],
            ["Registration window opens", "All eligible students", "In-app + SMS"],
            ["Course registration confirmed", "Student", "In-app + Email"],
            ["Results published", "Affected students", "In-app + SMS"],
            ["Payment received", "Student and bursar", "In-app + SMS"],
            ["Payment failed", "Student", "In-app + SMS"],
            ["Fee deadline in 7 days", "Students with a balance", "In-app + SMS"],
            ["Fee hold applied or removed", "Student", "In-app + SMS"],
            ["New assignment posted", "Enrolled students", "In-app"],
            ["Assignment graded", "Student", "In-app"],
            ["At-risk flag raised", "Lecturer and academic admin", "In-app + Email"],
            ["Timetable changed", "Affected students and lecturers", "In-app + SMS"],
            ["Account access revoked", "The user and the ICT manager", "Email"],
            ["Suspicious login", "The user and the ICT manager", "Email"],
        ],
    )

    h2(doc, "11.3 How it is built")
    bullets(
        doc,
        [
            "One notification service. Application code raises an event; it never calls the SMS "
            "provider directly.",
            "Templates live in the database so an ICT manager can edit wording without a "
            "deployment.",
            "Templates support English, French, and Kinyarwanda, matching the marketing site.",
            "Sending happens in a background queue with automatic retries, so a temporary "
            "provider outage does not lose messages.",
            "Bulk sends are throttled to respect provider limits.",
            "Users control their own preferences, except for legally required messages such as "
            "payment receipts.",
            "Delivery status is recorded, so a bursar can prove that a fee reminder was sent and "
            "delivered.",
        ],
    )

    # ── 12. Payments ─────────────────────────────────────────────────────────
    h1(doc, "12. Payments and Money Handling")

    para(
        doc,
        "This is the highest-risk area in the whole system. A bug in results annoys "
        "people; a bug in payments loses money and destroys trust. It gets the "
        "strictest rules and the most testing.",
    )

    h2(doc, "12.1 Payment flow for mobile money")
    numbered(
        doc,
        [
            "The student chooses MTN MoMo and enters their phone number.",
            "The browser calls POST /payments/initiate with an idempotency key, the invoice ID, "
            "and the method. It does not send the amount — the server looks that up itself.",
            "The server creates a payment record with status 'Pending' and calls the MoMo "
            "Collections API.",
            "MoMo sends a USSD push to the student's phone.",
            "The student approves with their PIN.",
            "MoMo calls our webhook at /webhooks/momo.",
            "We verify the webhook signature. An unsigned or badly signed webhook is rejected "
            "and logged as a security event.",
            "We confirm the amount matches the invoice exactly, then mark the payment "
            "'Completed'.",
            "A receipt PDF is generated in the background and stored in R2.",
            "The student and the bursar are notified.",
            "Meanwhile the browser polls the payment status so the screen updates by itself.",
        ],
    )

    h2(doc, "12.2 Non-negotiable rules")
    bullets(
        doc,
        [
            "Amounts always come from the server. Never trust an amount sent by the browser.",
            "Every payment attempt carries an idempotency key, so a student tapping 'Pay' three "
            "times results in one charge, not three.",
            "Follow-up work is queued in the same database transaction as the payment record. "
            "Marking a payment complete and queueing its receipt either both succeed or both "
            "fail. With an external queue service, the payment could be saved and the receipt "
            "job silently lost, leaving a student who has paid with no proof.",
            "Webhooks are verified by signature and are safe to receive more than once. Payment "
            "gateways do resend.",
            "If a webhook never arrives, a scheduled job checks pending payments against the "
            "gateway every 15 minutes and reconciles them.",
            "Payments are never deleted. A mistake is corrected with a reversal record that "
            "references the original.",
            "Receipt numbers run in a strict sequence per institution, with no gaps, because "
            "auditors check this.",
            "Every payment record links to the exact gateway reference for dispute resolution.",
            "All money is whole Rwandan Francs, stored as integers.",
        ],
    )

    h2(doc, "12.3 Reconciliation")
    bullets(
        doc,
        [
            "A daily job compares our records with each gateway's settlement report.",
            "Anything that does not match appears on the existing bursar reconciliation screen.",
            "Common cases handled: money received but no webhook, webhook received but money "
            "not settled, duplicate charge, wrong amount, payment made against the wrong "
            "student ID.",
            "The bursar can match a stray payment to a student manually, and that action is "
            "logged.",
            "Bank transfers are reconciled by importing a bank statement file, since banks do "
            "not send webhooks.",
        ],
    )

    h2(doc, "12.4 Fee holds")
    bullets(
        doc,
        [
            "A bursar can place a hold on a student account when fees are unpaid.",
            "A hold blocks course registration and can optionally block result viewing.",
            "The rule is enforced in the API, not only hidden in the interface. A student who "
            "calls the endpoint directly must still be blocked.",
            "The hold is removed automatically the moment full payment is received, so nobody "
            "has to remember to lift it.",
            "Applying and removing a hold is always logged with a reason.",
        ],
    )

    # ── 13. Phases ───────────────────────────────────────────────────────────
    h1(doc, "13. Implementation Phases")

    para(
        doc,
        "The work is divided into 18 phases across 5 stages. Phases 0 to 3 must be done "
        "in order, because everything else depends on them, so Stage 1 is the sum of its "
        "phases. Phases 4 to 10 build the role portals and overlap heavily once the "
        "foundation is stable, which is why Stage 2 is shorter than its phases added "
        "together. Those figures assume the team in section 14.1 working in parallel.",
    )

    table(
        doc,
        ["Stage", "Phases", "Focus", "Rough duration"],
        [
            ["Stage 1", "0 to 3", "Foundation: hosting, database, login, file storage", "11 weeks"],
            ["Stage 2", "4 to 10", "The seven role portals and the public admissions site", "20 weeks"],
            ["Stage 3", "11 to 13", "Notifications, AI, reports and background jobs", "8 weeks"],
            ["Stage 4", "14 to 17", "Security, testing, deployment, pilot go-live", "8 weeks"],
            ["Stage 5", "18", "Mobile application", "12 weeks (after launch)"],
        ],
    )

    h2(doc, "Stage 1 — Foundation")

    phase(
        doc,
        0,
        "Project Setup and Hosting Move",
        "Get the ground ready before anyone writes a feature",
        "2 weeks",
        [
            (
                "What we do",
                [
                    "First, measure real latency from Kigali to Frankfurt and to Cape Town, on "
                    "both MTN and Airtel mobile data. Confirm the region before creating "
                    "anything, because a Render region is permanent.",
                    "Move apps/api from Cloudflare Wrangler to a plain Node server running Hono.",
                    "Create the Render web service in Frankfurt and connect api.stackedu.rw.",
                    "Create the Vercel project for the frontend and connect app.stackedu.rw and "
                    "app.stackedu.africa.",
                    "Create the Neon project in Frankfurt, enable pgvector, and switch off "
                    "scale-to-zero for production.",
                    "Configure CORS on the API and set cookies on the .stackedu.rw domain, so "
                    "login works across the two hosts.",
                    "Create the Cloudflare account with the R2 bucket and Cloudflare Images.",
                    "Create the Render Key Value store in the same Frankfurt region.",
                    "Set up environment variables in both Render and Vercel, for preview, "
                    "staging, and production.",
                    "Set up GitHub Actions: type checking, linting, and tests on every pull request.",
                    "Add branch protection so nothing reaches main without review.",
                    "Install Sentry in both the web app and the API.",
                    "Start the NCSA data controller registration, because it takes up to 30 working days.",
                    "Start the merchant account applications with MTN, Airtel, and DPO, which "
                    "have the longest lead times in the whole project.",
                ],
            ),
            (
                "Why this matters",
                [
                    "A Render region cannot be changed afterwards. Measuring latency now avoids "
                    "being stuck in the wrong place with no way out except a full migration.",
                    "Getting cookies and CORS working across two hosts on day one prevents "
                    "mysterious login failures later, which are slow and frustrating to debug.",
                    "Starting the legal and payment registrations now means they finish long "
                    "before they can block go-live.",
                ],
            ),
            (
                "Done when",
                [
                    "Latency measurements from Kigali are recorded and the region is confirmed "
                    "in writing.",
                    "A pull request automatically deploys both a frontend preview and an API "
                    "preview.",
                    "The API health endpoint responds from Frankfurt.",
                    "The frontend can call the API and receive a session cookie that persists "
                    "across page reloads.",
                    "The API connects to Neon using the pooled connection string.",
                    "The NCSA application has been submitted.",
                ],
            ),
        ],
    )

    phase(
        doc,
        1,
        "Database Schema and Shared Contracts",
        "Define the data once, and build the tools that manage many databases",
        "4 weeks",
        [
            (
                "What we do",
                [
                    "Write the small platform schema: institutions, the email-to-institution "
                    "directory, and platform audit.",
                    "Write the full Drizzle schema for an institution database, covering all the "
                    "tables listed in section 5.1.",
                    "Build the institution provisioning script that creates a new Neon database, "
                    "applies the schema, and registers it in the platform database.",
                    "Build the migration runner that applies a change across every institution "
                    "database in turn, records progress, and stops safely on failure.",
                    "Build the connection manager that opens a pool per institution on first use "
                    "and closes it after a period of inactivity.",
                    "Write Zod validation schemas in packages/shared for every request and "
                    "response. This folder is currently empty and is the highest-leverage work in "
                    "the project.",
                    "Set up automatic OpenAPI generation from those Zod schemas.",
                    "Add indexes for the queries we know will be common: student lookup by "
                    "institution and programme, results by semester, payments by date.",
                    "Enable Row Level Security policies on every table.",
                    "Write a seed script that creates one demo institution with realistic data, "
                    "replacing the hard-coded mock files.",
                    "Set up automatic backups and run one restore test.",
                ],
            ),
            (
                "Why this order",
                [
                    "Every later phase reads and writes these tables. Changing the schema after "
                    "five portals are built is expensive; changing it now costs nothing.",
                    "The existing types in packages/shared/src/types/ are the starting point, so "
                    "this is refinement rather than invention.",
                ],
            ),
            (
                "Done when",
                [
                    "Migrations run cleanly on an empty database.",
                    "Provisioning a brand new institution takes one command and under five "
                    "minutes.",
                    "The migration runner applies a schema change across three test institutions "
                    "and reports the result for each.",
                    "The seed script produces two browsable demo institutions, so cross-tenant "
                    "mistakes are visible immediately during development.",
                    "The OpenAPI document generates without errors.",
                    "A test proves that a session for one institution cannot reach another "
                    "institution's database.",
                ],
            ),
        ],
    )

    phase(
        doc,
        2,
        "Authentication, Roles and Audit",
        "Nothing else can be trusted until this is right",
        "3 weeks",
        [
            (
                "What we do",
                [
                    "Install and configure Better Auth against our PostgreSQL database.",
                    "Build the endpoints: login, verify OTP, logout, forgot password, reset "
                    "password, session, refresh.",
                    "Wire up the four existing screens: login.tsx, verify.tsx, "
                    "forgot-password.tsx, reset-password.tsx.",
                    "Replace the stub in routes/_auth.tsx with a real session check.",
                    "Add role-based route guards, so a student cannot open a bursar screen by "
                    "typing the URL.",
                    "Implement useCurrentUser() properly — it currently always returns null.",
                    "Build the RBAC middleware and store the permission matrix from section 7.1 "
                    "in the database.",
                    "Build the audit log service and call it from every write operation.",
                    "Implement immediate session revocation.",
                    "Add two-factor authentication using the existing OTP input component.",
                ],
            ),
            (
                "Security focus",
                [
                    "Write tests that attempt to access every role's endpoints as every other "
                    "role. All should fail with 403.",
                    "Write tests that attempt to read another institution's data. All should fail.",
                ],
            ),
            (
                "Done when",
                [
                    "All six roles can log in and see only their own section.",
                    "Two-factor authentication works by SMS and email.",
                    "Revoking a user ends their session within seconds.",
                    "Every write action appears in the audit log with who, what, and when.",
                ],
            ),
        ],
    )

    phase(
        doc,
        3,
        "File Storage and Upload Pipeline",
        "Build it once, use it in six modules",
        "2 weeks",
        [
            (
                "What we do",
                [
                    "Build the three file endpoints: request upload link, confirm upload, "
                    "request download link.",
                    "Connect Cloudflare R2 with pre-signed URLs.",
                    "Connect Cloudflare Images for pictures, with automatic resizing and format "
                    "conversion.",
                    "Build a reusable FileUpload React component with drag and drop, progress "
                    "bar, and clear error messages. The specification lists this as a component "
                    "that must be built new.",
                    "Add virus scanning as a background job.",
                    "Add file type verification that checks the actual content, not just the "
                    "file extension.",
                    "Add the cleanup job for abandoned uploads.",
                    "Serve files from a separate domain.",
                ],
            ),
            (
                "Why now, before the portals",
                [
                    "Admissions, assignments, course materials, and the e-library all upload "
                    "files. Building this once prevents four different half-working "
                    "implementations.",
                ],
            ),
            (
                "Done when",
                [
                    "A 200 MB PDF uploads successfully from a browser on a normal connection.",
                    "An unauthorised user cannot obtain a download link.",
                    "A file with a fake extension is rejected.",
                    "Images are automatically served smaller on mobile.",
                ],
            ),
        ],
    )

    h2(doc, "Stage 2 — The Role Portals")
    para(
        doc,
        "Admissions, Student, ICT Manager, and Academic Admin are live on real data. "
        "Bursar is next so fee holds can block registration correctly. Every portal page "
        "includes a short, dismissible guide that explains what the signed-in person can "
        "do on that screen.",
    )

    phase(
        doc,
        4,
        "Admissions and the Public Application Portal",
        "7 screens · LIVE on real records · email verification before sign-in",
        "Done",
        [
            (
                "Screens covered",
                [
                    "apply/index (create account + verification modal), apply/form (7 steps), "
                    "apply/documents, apply/payment, apply/verify (returning unverified applicants), "
                    "apply/confirmation, apply/track",
                    "academic/applications (inbox) and academic/application (detail and decision)",
                ],
            ),
            (
                "Backend work (live)",
                [
                    "POST /apply/register creates the applicant account and draft application, "
                    "sends a six-digit email code, and does not issue a session.",
                    "POST /apply/verify-email and POST /apply/resend-verification are public; "
                    "verify checks the code, marks emailVerifiedAt, then signs the applicant in.",
                    "Unverified applicants cannot sign in or call form, document, or payment APIs.",
                    "Application create, save-as-draft, and submit endpoints.",
                    "Document upload through presigned R2 or local storage.",
                    "Application fee payment in sandbox mode (ICT toggles gate live MoMo/Airtel).",
                    "Review workflow: submitted, under review, documents requested, accepted, "
                    "rejected, with Resend email when the integration is on.",
                ],
            ),
            (
                "Institution branding (live)",
                [
                    "ICT System Settings: institution name, short name, contact email, website URL, "
                    "location, school logo upload, timezone, and locale.",
                    "GET /public/institution/:slug returns name, shortName, website, location, "
                    "and logoUrl for login, apply, and verification screens.",
                ],
            ),
            (
                "Done when",
                [
                    "An applicant registers, verifies email in the modal, and only then reaches "
                    "the application form.",
                    "An applicant who closes the browser before verifying can complete verification "
                    "at /apply/verify with email, password, and code — still without accessing the "
                    "form until verified.",
                    "An academic admin can review and decide on submitted applications.",
                ],
            ),
        ],
    )

    phase(
        doc,
        5,
        "Student Portal",
        "15 screens · LIVE on real records · AI paused",
        "Done",
        [
            (
                "Screens covered",
                [
                    "dashboard, onboarding, courses, course-detail, course-registration, "
                    "assignment-submit, results, transcript, timetable, fees, payment, receipt, "
                    "library, notifications, index",
                ],
            ),
            (
                "Backend work",
                [
                    "Student profile and dashboard summary endpoints.",
                    "Onboarding checklist with saved progress.",
                    "Course registration with credit limits, prerequisite checks, registration "
                    "window checks, and fee hold checks.",
                    "Enrolled courses, materials, and announcements.",
                    "Assignment submission with file upload.",
                    "Results and GPA/CGPA calculation, plus transcript PDF generation.",
                    "Personal timetable.",
                    "Fee statement, payment initiation, and receipt download.",
                    "E-library browsing and download, respecting access rules.",
                    "Notification list with read state.",
                ],
            ),
            (
                "Careful points",
                [
                    "The GPA and CGPA engine must handle repeated courses, resits, credit "
                    "weighting, and the institution's own grading scale. This is deterministic "
                    "code with heavy unit tests, not AI.",
                    "Registration must reject invalid attempts on the server, even if the "
                    "interface already hides the option.",
                    "Most students are on mobile data, so every screen must load quickly on a "
                    "slow connection.",
                ],
            ),
            (
                "Done when",
                [
                    "A student can register for courses, submit an assignment, view results, "
                    "download a transcript, pay fees, and open a library book.",
                    "A student cannot see any other student's data, verified by test.",
                ],
            ),
        ],
    )

    phase(
        doc,
        6,
        "Lecturer Portal",
        "14 screens · the staff who use it most often",
        "3 weeks",
        [
            (
                "Screens covered",
                [
                    "dashboard, courses, course-management, attendance, attendance-history, "
                    "results, result-review, assignments, submission-review, "
                    "assessment-builder, analytics, at-risk, notifications, index",
                ],
            ),
            (
                "Backend work",
                [
                    "Assigned courses and class rosters, scoped strictly to the lecturer's own "
                    "courses.",
                    "Attendance sessions and per-student records, with an export.",
                    "Material upload and course announcements.",
                    "Assignment creation, submission listing, grading with feedback.",
                    "Result entry per assessment component, saved as draft then submitted.",
                    "Online assessment builder with a question bank and automatic grading for "
                    "objective questions.",
                    "Course analytics: grade distribution and attendance trends.",
                    "At-risk list, reading from the rules engine.",
                ],
            ),
            (
                "Security focus",
                [
                    "Every single endpoint must confirm the lecturer is assigned to the course "
                    "in question. This is the most likely place for a permission bug.",
                    "Once results are published, a lecturer can no longer edit them. Only an "
                    "academic admin can correct them, and the correction is audited.",
                ],
            ),
            (
                "Done when",
                [
                    "A lecturer can take attendance, upload material, set and grade an "
                    "assignment, and submit results for publishing.",
                    "A lecturer cannot touch a course they are not assigned to.",
                ],
            ),
        ],
    )

    phase(
        doc,
        7,
        "Academic Admin (Registrar) Portal",
        "18 screens · LIVE on real records · AI paused",
        "Done (timetable write and full faculty CRUD remain)",
        [
            (
                "Screens covered",
                [
                    "dashboard, applications, application, students, student, courses, "
                    "programmes, programme, calendar, timetable, faculty, results, reports, "
                    "at-risk, notifications, profile, settings, index",
                ],
            ),
            (
                "Backend work (live)",
                [
                    "Student registry with search, filters, and student detail views.",
                    "Applications inbox wired to the admissions API (review and decision).",
                    "Programme and course catalogue: create and edit via API.",
                    "Academic calendar: create, edit, and delete years, semesters, and events.",
                    "Timetable read from enrolled offerings (display only; no write API yet).",
                    "Faculty list from lecturer records (read-only; no add-or-assign API yet).",
                    "Result approval workflow (approve and reject with audit).",
                    "At-risk student list and academic reports from live aggregates.",
                    "In-app notifications list with read state.",
                    "Shared account profile and settings (notification preferences in DB).",
                ],
            ),
            (
                "Still to build",
                [
                    "Timetable manager with conflict detection and write endpoints.",
                    "Faculty management and course assignment to lecturers.",
                    "Enrolment status changes: suspend, transfer, graduate, withdraw, each with "
                    "an audit entry.",
                    "Result batch publishing as a background job with student notifications.",
                    "Result override with a mandatory reason and full audit trail.",
                    "Academic report export to Excel and PDF.",
                ],
            ),
            (
                "Done when",
                [
                    "A registrar can set the calendar, manage programmes and courses, review "
                    "applications, approve results, and monitor at-risk students — all on live data.",
                    "Timetable editing and full faculty assignment complete the remaining registrar "
                    "workflow.",
                    "An academic admin cannot see any payment data, verified by test.",
                ],
            ),
        ],
    )

    phase(
        doc,
        8,
        "Bursar Portal and Live Payments",
        "9 screens · the highest-risk module",
        "4 weeks",
        [
            (
                "Screens covered",
                [
                    "dashboard, fee-structure, ledger, student-account, student-accounts, "
                    "receipts, reports, reconciliation, index",
                ],
            ),
            (
                "Backend work",
                [
                    "Fee structure configuration by programme, year group, and student category.",
                    "Invoice generation for each student each semester.",
                    "Live integration with MTN MoMo Collections, Airtel Money, and DPO Pay.",
                    "Bank transfer recording and statement import.",
                    "Webhook handlers with signature verification and safe repeat handling.",
                    "The scheduled reconciliation job.",
                    "Receipt generation, reprinting, and voiding, with a strict number sequence.",
                    "Fee holds, enforced in the API.",
                    "Financial reports by date range, programme, channel, and cohort.",
                ],
            ),
            (
                "Testing focus",
                [
                    "Every gateway is tested in sandbox first, covering success, failure, "
                    "timeout, duplicate webhook, and wrong amount.",
                    "A specific test proves that pressing Pay three times produces one charge.",
                    "Reconciliation is tested against a deliberately mismatched dataset.",
                ],
            ),
            (
                "Done when",
                [
                    "A real payment of 100 RWF completes end to end on each channel in "
                    "production.",
                    "The ledger balances exactly against the gateway settlement report.",
                    "A bursar cannot see any academic results, verified by test.",
                ],
            ),
        ],
    )

    phase(
        doc,
        9,
        "Librarian Portal and E-Library",
        "7 screens · 6 of them must be built from nothing",
        "4 weeks",
        [
            (
                "The situation",
                [
                    "This is the only portal that is not designed or built. "
                    "routes/_auth/librarian/index.tsx contains a single line saying 'coming "
                    "soon'.",
                    "Unlike every other phase, this needs design work as well as development. "
                    "Budget for that.",
                    "The good news is that the student-facing library screen "
                    "(student/library.tsx, 1,798 lines) is fully built, so we already know what "
                    "the data must look like.",
                ],
            ),
            (
                "Screens to build",
                [
                    "Library Dashboard — totals, recently added, most accessed, pending requests",
                    "Resource Catalogue — searchable and filterable list of everything",
                    "Add Resource — metadata form plus file upload",
                    "Edit Resource — update metadata, replace file, change access rules",
                    "Collections Manager — thematic collections and course reading lists",
                    "Resource Requests — handle requests from students and lecturers",
                    "Library Analytics — most accessed, active users, search terms, downloads",
                ],
            ),
            (
                "Backend work",
                [
                    "Full resource management with metadata: title, author, ISBN, year, type, "
                    "subject tags.",
                    "Access rules by programme, department, and year group, enforced on every "
                    "download.",
                    "Collections and reading lists.",
                    "Access logging for every open and download, which feeds analytics.",
                    "Bulk import so an institution can load an existing catalogue.",
                    "Full-text search now; semantic search is added in Phase 12.",
                ],
            ),
            (
                "Done when",
                [
                    "A librarian can add a 200 MB e-book, restrict it to third-year law "
                    "students, and confirm that nobody else can download it.",
                    "The analytics screen shows real usage figures.",
                    "A librarian cannot see student records or payments, verified by test.",
                ],
            ),
        ],
    )

    phase(
        doc,
        10,
        "ICT Manager Portal",
        "15 screens · LIVE on real records · platform controls only",
        "Done",
        [
            (
                "Screens covered",
                [
                    "dashboard, users, user, access-levels, revocation, revocation-detail, "
                    "audit-log, audit-entry, settings, integrations, integration-detail, "
                    "analytics, announcements, notifications, profile, account-settings, index",
                ],
            ),
            (
                "What ICT owns (live)",
                [
                    "Create, edit, deactivate and restore logins for every role.",
                    "When the role is Student, ICT also creates the student record so the "
                    "person can sign in to the student portal.",
                    "Access levels: grant or remove the seeded permission keys per role.",
                    "Access revocation with a mandatory reason; all sessions for that user end.",
                    "Append-only audit log with human-readable labels. Even ICT cannot edit or "
                    "delete an entry.",
                    "Institution identity settings: name, short name, contact email, website URL, "
                    "location, school logo (upload to R2 or local storage), timezone, and locale. "
                    "Not faculties, programmes, or the academic calendar.",
                    "Integration on/off flags and connection health checks. When Resend is off, "
                    "outbound email is skipped; when MoMo or Airtel is off, live payments for that "
                    "channel are blocked. Secrets stay in environment variables, never in the "
                    "database.",
                    "Analytics dashboard from live audit and user activity aggregates.",
                    "Institution-wide announcements and ICT's own in-app notifications.",
                    "Shared account profile and settings (password change revokes all sessions; "
                    "TOTP two-factor authentication with QR setup; notification preferences in DB; "
                    "dark/light mode).",
                    "A short guide box on every ICT and Student page.",
                ],
            ),
            (
                "What ICT does not own",
                [
                    "Opening or closing a semester, registration windows, or 'current year'. "
                    "That is Academic Admin (calendar.write).",
                    "Fee structures, invoices, and fee holds. That is the Bursar.",
                    "Course materials, attendance, and marks. That is the Lecturer.",
                    "Publishing library resources. That is the Librarian.",
                ],
            ),
            (
                "Done when",
                [
                    "An ICT manager can create every type of user, change access levels, revoke "
                    "access instantly, and trace any action in the audit log.",
                    "Semester activation is not possible from the ICT portal.",
                ],
            ),
        ],
    )

    h2(doc, "Stage 3 — Intelligence and Platform Services")

    phase(
        doc,
        11,
        "Notifications",
        "The layer that makes the system feel alive",
        "2 weeks",
        [
            (
                "What we do",
                [
                    "Build the central notification service and the event triggers listed in "
                    "section 11.2.",
                    "Integrate Pindo for SMS with Africa's Talking as a fallback.",
                    "Integrate Resend for email with branded templates.",
                    "Build editable templates in three languages.",
                    "Build the in-app notification centre, connecting the existing notification "
                    "screens for each role.",
                    "Implement useNotifications(), which is currently a stub.",
                    "Add user preferences and delivery status tracking.",
                    "Add throttling and cost reporting for SMS.",
                ],
            ),
            (
                "Done when",
                [
                    "Publishing results sends SMS to exactly the right students and nobody else.",
                    "A failed SMS is retried and the failure is visible to an admin.",
                    "The monthly SMS cost per institution is reportable.",
                ],
            ),
        ],
    )

    phase(
        doc,
        12,
        "AI Layer",
        "The features that make it an AI-powered system",
        "4 weeks",
        [
            (
                "What we do",
                [
                    "At-risk detection rules engine, with explanations, feeding the existing "
                    "at-risk screens for lecturers and academic admins.",
                    "Intervention tracking so a follow-up can be assigned and closed.",
                    "Semantic library search using pgvector, with an embedding job for every "
                    "resource.",
                    "Timetable generation using a constraint solver, feeding the existing "
                    "timetable manager screen.",
                    "Document extraction during admissions review.",
                    "AI usage logging and per-institution cost caps.",
                    "Stretch item: the student study assistant, only if the schedule allows.",
                ],
            ),
            (
                "Deliberately not done yet",
                [
                    "A trained machine learning risk model. There is no historical outcome data "
                    "yet. It is added after two semesters of real use, and only then will it "
                    "beat the rules.",
                ],
            ),
            (
                "Done when",
                [
                    "At-risk students are flagged with visible, understandable reasons.",
                    "A search for a concept finds relevant books that do not contain that exact "
                    "word.",
                    "A conflict-free timetable is generated for a real programme.",
                    "AI cost per institution is visible and capped.",
                ],
            ),
        ],
    )

    phase(
        doc,
        13,
        "Reports, Documents and Background Jobs",
        "The heavy lifting that must never block a web request",
        "2 weeks",
        [
            (
                "What we do",
                [
                    "Set up Graphile Worker, a job queue that lives in our own PostgreSQL, run "
                    "by a Render background worker. Render's cron feature triggers the scheduled "
                    "jobs.",
                    "Because the queue is in the same database as the data, a job can be queued "
                    "inside the same transaction as the write that caused it. Either both happen "
                    "or neither does. Section 12.2 explains why this matters for payments.",
                    "Queued jobs are backed up along with the rest of the database, so a restore "
                    "brings back pending work as well as data.",
                    "Transcript PDF generation, matching the institution's official format.",
                    "Receipt PDF generation.",
                    "Report exports to Excel and PDF for academic, financial, library, and "
                    "platform reports.",
                    "Bulk operations: batch result publishing, bulk user import, bulk "
                    "notifications.",
                    "Scheduled jobs: nightly GPA recalculation, daily reconciliation, daily "
                    "at-risk scoring, weekly backup export, hourly abandoned-upload cleanup.",
                    "A job monitoring view so a failed job is visible and can be retried.",
                ],
            ),
            (
                "Why it is separate",
                [
                    "Generating 3,000 transcripts cannot happen inside a web request. It would "
                    "time out and leave the data half-written.",
                    "Jobs must be safe to retry, so a partial failure does not create duplicate "
                    "receipts.",
                ],
            ),
            (
                "Done when",
                [
                    "A transcript PDF is generated and downloaded within 10 seconds.",
                    "Publishing results for 2,000 students completes reliably in the background.",
                    "A deliberately failed job is retried automatically and reported.",
                ],
            ),
        ],
    )

    h2(doc, "Stage 4 — Hardening and Launch")

    phase(
        doc,
        14,
        "Security Hardening and Compliance",
        "Prove it is safe before real students depend on it",
        "2 weeks",
        [
            (
                "What we do",
                [
                    "Full review of every endpoint against the permission matrix.",
                    "Automated tests that attempt every role against every endpoint.",
                    "Automated tests that attempt cross-institution access on every table.",
                    "Add security headers: Content Security Policy, HSTS, and related headers.",
                    "Enable the Vercel firewall for the frontend and put Cloudflare in front of "
                    "api.stackedu.rw, then finalise rate limits.",
                    "Add field-level encryption for national ID numbers and similar fields.",
                    "Complete the NCSA registration and appoint the Data Protection Officer.",
                    "Write the privacy notice and the data processing agreement template.",
                    "Build the data export and correction request feature.",
                    "Run an internal penetration test, then an external one.",
                    "Write the incident response and breach notification plan.",
                ],
            ),
            (
                "Done when",
                [
                    "Every finding rated medium or above from the penetration test is fixed.",
                    "The NCSA registration certificate has been issued.",
                    "Legal documents are signed and published.",
                ],
            ),
        ],
    )

    phase(
        doc,
        15,
        "Testing and Quality Assurance",
        "Find the problems before the institution does",
        "3 weeks",
        [
            (
                "Test layers",
                [
                    "Unit tests for all business rules, with the GPA engine, fee calculations, "
                    "and permission checks covered heavily.",
                    "Integration tests for every API endpoint against a real test database.",
                    "End-to-end tests with Playwright for the critical journeys: apply and get "
                    "admitted, register for courses, submit an assignment, publish results, pay "
                    "fees, borrow a library resource, revoke a user.",
                    "Load testing simulating 5,000 students registering in the same hour, which "
                    "is exactly what happens when a registration window opens.",
                    "Accessibility testing to WCAG 2.1 AA.",
                    "Testing on real low-end Android devices on a real MTN or Airtel connection, "
                    "not only on office wifi.",
                    "User acceptance testing with actual staff from the pilot institution.",
                ],
            ),
            (
                "Targets",
                [
                    "Above 80 percent test coverage on business logic.",
                    "Every critical journey covered end to end.",
                    "Dashboard loads in under 3 seconds on a 3G connection.",
                    "No critical or high-severity bugs open at launch.",
                ],
            ),
        ],
    )

    phase(
        doc,
        16,
        "Deployment, Domains and Monitoring",
        "Make it live, and make sure we know if it breaks",
        "1 week",
        [
            (
                "What we do",
                [
                    "Connect app.stackedu.rw and app.stackedu.africa in Vercel, and "
                    "api.stackedu.rw in Render, then issue SSL certificates for all three.",
                    "Confirm the app does NOT geo-redirect between the two domains, unlike the "
                    "marketing site, so sessions are never lost.",
                    "Verify that a session started on .rw survives, and that the API rejects "
                    "requests from any origin other than our two app domains.",
                    "Set Render's minimum instance count to at least one, so there is never a "
                    "cold start for a real user.",
                    "Configure the files subdomain for downloads.",
                    "Set up the staging environment.",
                    "Configure Sentry alerts and uptime monitoring with on-call notification.",
                    "Set up a status page.",
                    "Document the rollback procedure and practise it once.",
                    "Write the runbook: what to do when payments fail, when SMS fails, when the "
                    "database is slow.",
                ],
            ),
            (
                "Done when",
                [
                    "Both domains serve the app over HTTPS.",
                    "A deliberate error raises an alert within one minute.",
                    "A rollback has been performed successfully in a drill.",
                ],
            ),
        ],
    )

    phase(
        doc,
        17,
        "Pilot, Data Migration, Training and Go-Live",
        "One institution first, carefully",
        "4 weeks",
        [
            (
                "What we do",
                [
                    "Migrate the pilot institution's existing data using the process in "
                    "section 5.6.",
                    "Configure their branding, programmes, courses, fee structures, and academic "
                    "calendar.",
                    "Train the ICT manager first, then academic admins, bursars, librarians, and "
                    "lecturers.",
                    "Produce short video guides and one-page quick reference sheets in English "
                    "and Kinyarwanda.",
                    "Run student orientation sessions.",
                    "Go live with one department or one year group before the whole institution.",
                    "Provide daily support for the first two weeks, then weekly.",
                    "Collect feedback and fix issues quickly.",
                ],
            ),
            (
                "Why a limited go-live",
                [
                    "If something is wrong, it affects 200 students instead of 8,000.",
                    "Staff confidence grows before the pressure of a full registration period.",
                ],
            ),
            (
                "Done when",
                [
                    "The pilot group completes a full registration and payment cycle in the live "
                    "system.",
                    "Staff can operate the system without daily help.",
                    "The institution formally signs off.",
                ],
            ),
        ],
    )

    h2(doc, "Stage 5 — After Launch")

    phase(
        doc,
        18,
        "Mobile Application",
        "Android first, then iOS",
        "12 weeks",
        [
            (
                "What we do",
                [
                    "Add apps/mobile to the monorepo using React Native and Expo.",
                    "Reuse packages/shared for types and validation, so nothing is duplicated.",
                    "Build the student experience first, since students are the largest group and "
                    "are already mobile-first.",
                    "Add push notifications through Expo, which removes most SMS cost over time.",
                    "Add offline reading, so downloaded course materials and library resources "
                    "work without a connection.",
                    "Add biometric login.",
                    "Publish to Google Play first, then the Apple App Store.",
                    "Later releases add the lecturer experience: attendance and quick grade entry.",
                ],
            ),
            (
                "Why it is cheap by then",
                [
                    "The API already exists, is versioned, and already issues tokens.",
                    "Types, validation, and design tokens are already shared.",
                    "Only the interface layer is new.",
                ],
            ),
        ],
    )

    # ── 14. Team and timeline ────────────────────────────────────────────────
    h1(doc, "14. Team and Timeline")

    h2(doc, "14.1 Suggested team")
    table(
        doc,
        ["Role", "Number", "Main responsibility"],
        [
            ["Backend engineer", "2", "API, database, integrations, jobs"],
            ["Frontend engineer", "1 to 2", "Connect screens to the API, build the Librarian portal"],
            ["Full-stack engineer", "1", "Moves between both, covers gaps"],
            ["Product designer", "0.5", "Librarian portal design, refinements"],
            ["QA engineer", "0.5 rising to 1", "Test plans, automation, user acceptance testing"],
            ["DevOps and security", "0.5", "Pipelines, monitoring, compliance"],
            ["Project lead", "1", "Client relationship, scope, delivery"],
        ],
    )

    h2(doc, "14.2 Two possible plans")
    table(
        doc,
        ["Plan", "Scope", "Time to live"],
        [
            [
                "Minimum pilot",
                "Phases 0 to 8, plus 11, 15, 16, 17. Admissions, students, lecturers, academic "
                "admin, and payments. No library portal, no AI.",
                "About 6 months",
            ],
            [
                "Full system",
                "All phases 0 to 17. Everything including the Librarian portal and the AI layer.",
                "About 9 to 11 months",
            ],
        ],
    )
    para(
        doc,
        "The minimum pilot is the recommended path. It gets a real institution using "
        "the system and paying for it while the remaining modules are built, and real "
        "usage produces the data the AI features need anyway.",
    )

    # ── 15. Costs ────────────────────────────────────────────────────────────
    h1(doc, "15. Estimated Running Costs")

    para(
        doc,
        "These are indicative monthly figures in US dollars for one institution of "
        "roughly 5,000 students. Verify current prices with each provider before "
        "committing, as they change.",
    )

    table(
        doc,
        ["Service", "Development", "Live with one institution"],
        [
            ["Vercel (frontend only)", "0 to 20", "20 to 40"],
            ["Render web service", "0 (free tier)", "25 to 85"],
            ["Render worker and cron", "0", "7 to 25"],
            [
                "Neon PostgreSQL",
                "0 (free tier)",
                "35 to 180 — one database per institution plus the platform database, on a plan "
                "that allows several projects",
            ],
            ["Cloudflare R2", "under 5", "10 to 40 (no download charges)"],
            ["Cloudflare Images", "0 (free tier)", "5 to 25"],
            ["Render Key Value", "0", "10 to 30"],
            ["Job queue", "0", "0 — runs inside PostgreSQL, no separate vendor"],
            ["Resend email", "0", "10 to 30"],
            ["SMS (Pindo)", "5", "50 to 300, depends heavily on volume"],
            ["Sentry", "0", "26"],
            ["Uptime monitoring", "0", "10"],
            ["AI model usage", "10", "30 to 150"],
            ["Estimated total", "about 40", "about 230 to 1,000"],
        ],
    )

    bullets(
        doc,
        [
            "SMS is the most variable cost and the easiest to overspend. Every template must be "
            "justified, and the mobile app with push notifications will reduce it substantially.",
            "Choosing Cloudflare R2 over a bandwidth-billed service is the single largest saving "
            "in this table once the e-library is in real use.",
            "Render's fixed monthly pricing makes the total far easier to predict than "
            "usage-based serverless billing, which matters when quoting a fixed annual fee to an "
            "institution.",
            "Neon's autoscaling means we pay for the registration-day peak only on the days it "
            "actually happens, rather than all year round.",
            "One database per institution costs somewhat more than a single shared database, but "
            "the amount is small compared with the price of a bad restore, and it scales in step "
            "with revenue because each new database arrives with a new paying institution.",
            "Dropping Cloudinary, Upstash, and a separate job-queue service removes three "
            "vendors, three bills, and three sets of credentials to manage.",
            "Costs per institution fall as more institutions join, because most of these "
            "services are shared across all of them.",
        ],
    )

    # ── 16. Risks ────────────────────────────────────────────────────────────
    h1(doc, "16. Risks and How We Manage Them")

    table(
        doc,
        ["Risk", "Impact", "How we reduce it"],
        [
            [
                "Payment integration takes longer than planned",
                "High — it blocks the whole finance module",
                "Start gateway account applications in Phase 0, long before Phase 8. Approval "
                "with MTN and Airtel can take weeks.",
            ],
            [
                "The Librarian portal is undesigned",
                "Medium — it is a whole missing module",
                "Design work starts during Phase 7 so development is not waiting. Budget "
                "explicitly for it.",
            ],
            [
                "NCSA registration is delayed",
                "High — it can legally block go-live",
                "Submit in Phase 0. It takes up to 30 working days and is free.",
            ],
            [
                "Poor connectivity for students",
                "High — students cannot use the system",
                "Small page sizes, aggressive caching, offline-tolerant design, and testing on "
                "real mobile networks rather than office wifi.",
            ],
            [
                "Migrating messy institutional data",
                "High — bad data destroys trust immediately",
                "Migrate into staging first, have the institution verify a sample, and run "
                "automatic reconciliation checks before production.",
            ],
            [
                "Registration day overload",
                "High — everyone registers in the same hour",
                "A long-lived Render server with a fixed connection pool, plus Neon's "
                "autoscaling compute, handles this far more predictably than serverless would. "
                "Load test for 5,000 concurrent users, use queues, and stagger registration "
                "windows by year group.",
            ],
            [
                "Frankfurt proves too slow for Rwandan users",
                "Medium — the app feels sluggish on every screen",
                "Measure from Kigali in Phase 0 before committing, because a Render region is "
                "permanent. If Frankfurt tests badly, fall back to Vercel Cape Town with "
                "Supabase Cape Town, which stays available until Phase 0 closes.",
            ],
            [
                "A migration succeeds on some institutions and fails on others",
                "Medium — the platform is left in an inconsistent state",
                "The migration runner records progress per institution and stops on the first "
                "failure rather than continuing. Every migration is rehearsed on a Neon branch "
                "first. Changes are written so that old and new code both work during a rollout.",
            ],
            [
                "Login breaks between the two hosts",
                "Medium — users cannot stay signed in",
                "The frontend and API share the stackedu.rw domain, so cookies work with "
                "Domain=.stackedu.rw. Configure and test this in Phase 0, not in Phase 2 when it "
                "would block everything.",
            ],
            [
                "Staff resistance to a new system",
                "Medium — the system is ignored",
                "Involve staff during user acceptance testing, train early, and start with one "
                "department.",
            ],
            [
                "Scope grows during the build",
                "Medium — the launch date slips",
                "This document is the agreed scope. Anything new goes into a version two list "
                "unless it is formally accepted with a new date.",
            ],
            [
                "Depending on a single AI provider",
                "Low to medium",
                "Use the Vercel AI SDK, which allows switching provider without rewriting, and "
                "make sure no critical function requires AI.",
            ],
        ],
    )

    # ── 17. Next steps ───────────────────────────────────────────────────────
    h1(doc, "17. Immediate Next Steps")

    para(doc, "These are the decisions and actions needed before development starts.")

    numbered(
        doc,
        [
            "Approve the technology stack in section 3.8: TypeScript with Hono on Render, Neon "
            "for the database with one database per institution, and Cloudflare R2 with "
            "Cloudflare Images for files.",
            "Choose the delivery plan: minimum pilot in about 5 months, or full system in about "
            "9 to 11 months.",
            "Run the Kigali latency test against Frankfurt and Cape Town, and confirm the region "
            "in writing before any Render service is created. This is the one decision that "
            "cannot be undone cheaply.",
            "Create the Render and Neon accounts, both in Frankfurt.",
            "Begin the NCSA data controller registration today, since it takes up to 30 working "
            "days and can block go-live.",
            "Start the merchant account applications with MTN MoMo, Airtel Money, and DPO Pay. "
            "These have the longest lead times of anything in the project.",
            "Confirm the pilot institution and get access to a sample of their existing data.",
            "Commission the design work for the seven Librarian screens.",
            "Confirm the team and start Phase 0.",
        ],
    )

    para(
        doc,
        "Much of this plan is already built and live on Render and Vercel for the pilot "
        "institution. Remaining work is concentrated in Bursar, Lecturer, Librarian, live "
        "payment webhooks, and the central notification service. Regenerate this document "
        "locally with scripts/generate-implementation-plan.py whenever a major milestone ships.",
    )

    # ── Appendix ─────────────────────────────────────────────────────────────
    h1(doc, "Appendix A — Full Screen Inventory")

    para(
        doc,
        "Every screen file found in apps/web/src/routes, grouped by area, with the "
        "phase in which it is connected to real data.",
    )

    table(
        doc,
        ["Area", "Screens", "Phase"],
        [
            ["Public", "index, login, verify, forgot-password, reset-password, demo", "2"],
            [
                "Apply",
                "index, form, documents, payment, verify, confirmation, track (live; verify "
                "before sign-in)",
                "4 — Done",
            ],
            [
                "Student",
                "index, dashboard, onboarding, courses, course-detail, course-registration, "
                "assignment-submit, results, transcript, timetable, fees, payment, receipt, "
                "library, notifications",
                "5",
            ],
            [
                "Lecturer",
                "index, dashboard, courses, course-management, attendance, attendance-history, "
                "results, result-review, assignments, submission-review, assessment-builder, "
                "analytics, at-risk, notifications",
                "6",
            ],
            [
                "Academic Admin",
                "index, dashboard, applications, application, students, student, courses, "
                "programmes, programme, calendar, timetable, faculty, results, reports, at-risk, "
                "notifications",
                "7",
            ],
            [
                "Bursar",
                "index, dashboard, fee-structure, ledger, student-account, student-accounts, "
                "receipts, reports, reconciliation",
                "8",
            ],
            [
                "Librarian",
                "index (placeholder only). To build: dashboard, catalogue, add-resource, "
                "edit-resource, collections, requests, analytics",
                "9",
            ],
            [
                "ICT Manager",
                "index, dashboard, users, user, access-levels, revocation, revocation-detail, "
                "audit-log, audit-entry, settings, integrations, integration-detail, analytics, "
                "announcements, notifications",
                "10",
            ],
        ],
    )

    h1(doc, "Appendix B — Files That Must Be Replaced")

    para(
        doc,
        "These files currently hold fake data or placeholder logic. Each one is "
        "replaced with real API calls during the phase shown.",
    )

    table(
        doc,
        ["File", "What it is now", "Phase"],
        [
            ["apps/api/src/index.ts", "One health endpoint only", "0 to 1"],
            ["packages/shared/src/schemas/index.ts", "Empty file", "1"],
            ["apps/web/src/routes/_auth.tsx", "Contains 'isAuthenticated = true'", "2"],
            ["apps/web/src/hooks/useCurrentUser.ts", "Always returns null", "2"],
            ["apps/web/src/hooks/useNotifications.ts", "Stub", "11"],
            ["apps/web/src/lib/api/students.ts", "Throws 'not implemented'", "5 and 7"],
            ["apps/web/src/lib/api/courses.ts", "Throws 'not implemented'", "5 to 7"],
            ["apps/web/src/lib/api/fees.ts", "Throws 'not implemented'", "8"],
            ["apps/web/src/lib/api/results.ts", "Throws 'not implemented'", "5 to 7"],
            ["apps/web/src/lib/api/library.ts", "Throws 'not implemented'", "9"],
            ["apps/web/src/data/academic.ts", "1,144 lines of fake data", "7"],
            ["apps/web/src/data/ict.ts", "847 lines of fake data", "10"],
            ["apps/web/src/data/courses.ts", "684 lines of fake data", "5 to 7"],
            ["apps/web/src/data/lecturer.ts", "422 lines of fake data", "6"],
            ["apps/web/src/data/bursar.ts", "398 lines of fake data", "8"],
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
