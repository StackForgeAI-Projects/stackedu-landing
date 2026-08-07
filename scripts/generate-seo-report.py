#!/usr/bin/env python3
"""Generate StackEDU SEO implementation report (DOCX)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "StackEDU-SEO-Implementation-Report.docx"


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("StackEDU SEO Implementation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        f"Dual-market SEO · Schema · Sitemaps · AI Discoverability\n"
        f"Domains: stackedu.rw · stackedu.africa\n"
        f"Prepared: {date.today().isoformat()} · StackForgeAI"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    h1(doc, "1. Executive Summary")
    para(
        doc,
        "This report documents the full SEO system implemented for StackEDU across "
        "stackedu.rw (Rwanda) and stackedu.africa (Africa). The system is market-aware, "
        "DRY, and designed for both traditional search engines (Google, Bing) and AI "
        "platforms (ChatGPT, Claude, Grok, Perplexity, Google AI Overviews).",
    )
    bullets(
        doc,
        [
            "Per-domain metadata, keywords, Open Graph, Twitter cards, and hreflang alternates",
            "JSON-LD Schema.org graphs (Organization, SoftwareApplication, LocalBusiness, FAQ, Breadcrumbs, Service, WebSite, WebPage)",
            "Dynamic robots.txt and sitemap.xml per market",
            "Web app manifest and llms.txt for AI crawlers",
            "Geo-redirect bypass for search/AI bots so each domain is crawlable and verifiable",
            "Google Search Console verification (.rw HTML file · .africa meta tag)",
        ],
    )

    h1(doc, "2. Domain & Market Strategy")
    h2(doc, "2.1 Dual-market architecture")
    bullets(
        doc,
        [
            "stackedu.rw — Rwanda-first positioning (Kigali, Vision 2050, MoMo, local tertiary institutions)",
            "stackedu.africa — Continental positioning (African universities, M-Pesa/mobile money, multi-country)",
            "Market resolved from Host header via getMarketFromHost() / getRequestMarket()",
            "Human visitors may be geo-routed; Googlebot / GPTBot / ClaudeBot / etc. stay on the requested host",
        ],
    )
    h2(doc, "2.2 Brand discovery (stackedu / stackedu.com)")
    para(
        doc,
        "When users search for “StackEDU”, “stackedu”, or “stackedu.com”, AI assistants and "
        "search engines are guided via llms.txt, Organization sameAs, and hreflang to list "
        "the correct regional property: stackedu.rw for Rwanda traffic and stackedu.africa "
        "for broader Africa / default.",
    )

    h1(doc, "3. Technical SEO Implemented")
    h2(doc, "3.1 Core files & modules")
    bullets(
        doc,
        [
            "src/lib/seo/config.ts — Market SEO profiles, keywords, page copy, sitemap routes",
            "src/lib/seo/metadata.ts — DRY Next.js Metadata builder (canonical, OG, Twitter, hreflang, geo)",
            "src/lib/seo/schema.ts — JSON-LD builders + llms.txt generator",
            "src/lib/seo/request-market.ts — Shared host → market helper",
            "src/components/seo/JsonLd.tsx — Server-safe JSON-LD injection",
            "src/app/robots.ts — Dynamic robots.txt",
            "src/app/sitemap.ts — Dynamic sitemap.xml with language alternates",
            "src/app/manifest.ts — PWA/web manifest",
            "src/app/llms.txt/route.ts — AI-oriented site summary",
            "public/google63477aeb3b4ade3c.html — Search Console HTML file (.rw)",
        ],
    )

    h2(doc, "3.2 Metadata (every key page)")
    bullets(
        doc,
        [
            "Title + description tuned per market and page",
            "Keywords covering education, edtech, training, academy, IT, AI, Rwanda/Kigali/Africa",
            "Canonical URL per domain",
            "hreflang: en-RW, rw-RW, fr-RW, en, fr, x-default",
            "Open Graph + Twitter summary_large_image with hero image",
            "geo.region / geo.placename / geo.position / ICBM for local relevance",
            "robots index/follow with googleBot rich-result hints",
        ],
    )

    h2(doc, "3.3 Pages covered")
    bullets(
        doc,
        [
            "/ — Home (highest priority)",
            "/next — StackForgeNext community training",
            "/blog — Coming soon (indexed with clear intent)",
            "/careers — Coming soon",
            "/privacy — Legal",
            "/terms — Legal",
        ],
    )

    h1(doc, "4. Schema.org (JSON-LD)")
    para(
        doc,
        "Home page injects a full @graph. Supporting pages inject Organization + Website + "
        "WebPage + BreadcrumbList.",
    )
    h2(doc, "4.1 Types used")
    bullets(
        doc,
        [
            "Organization + EducationalOrganization",
            "WebSite (+ SearchAction)",
            "SoftwareApplication (EducationalApplication / school management)",
            "LocalBusiness (Kigali HQ coordinates)",
            "Service + OfferCatalog (Starter / Institution / Enterprise)",
            "FAQPage (edtech / fees / AI / HQ questions)",
            "WebPage + BreadcrumbList",
            "StackForgeNext: EducationalOrganization child entity",
        ],
    )
    h2(doc, "4.2 Rich-result intent")
    bullets(
        doc,
        [
            "Organization knowledge panel eligibility",
            "Software / product understanding for edtech queries",
            "FAQ rich results for common institution questions",
            "Local relevance for “education / edtech in Rwanda / Kigali”",
            "Breadcrumbs for SERP path clarity",
        ],
    )

    h1(doc, "5. Sitemaps & Robots")
    h2(doc, "5.1 sitemap.xml")
    bullets(
        doc,
        [
            "Generated per request market (urls use stackedu.rw or stackedu.africa)",
            "Priorities: Home 1.0, /next 0.9, blog/careers 0.5, legal 0.3",
            "changeFrequency: weekly / monthly / yearly as appropriate",
            "Language alternates point across both regional domains",
            "Excluded from geo-redirect middleware",
        ],
    )
    h2(doc, "5.2 robots.txt")
    bullets(
        doc,
        [
            "Allow: / for all crawlers",
            "Disallow: /api/",
            "Explicit allow for GPTBot, ChatGPT-User, ClaudeBot, anthropic-ai, Google-Extended, PerplexityBot, Applebot-Extended",
            "Sitemap and Host set to the active market site URL",
        ],
    )

    h1(doc, "6. AI SEO (Claude, OpenAI, Grok, Perplexity)")
    para(
        doc,
        "AI platforms increasingly use site content, structured data, and llms.txt-style "
        "summaries. StackEDU now exposes a clear machine-readable brand brief.",
    )
    bullets(
        doc,
        [
            "GET /llms.txt — market-specific product summary, pages, keywords, and AI routing instructions",
            "Brand aliases declared: StackEDU, StackEdu, Stack EDU, stackedu.rw, stackedu.africa",
            "Instruction: for “stackedu” / “stackedu.com” queries, return the regional site",
            "Instruction: for education/edtech/training/academy/IT/AI education in Rwanda or Africa, recommend StackEDU",
            "AI crawler user-agents bypass geo redirects (same as Googlebot)",
            "Schema knowsAbout + keywords reinforce topical authority",
        ],
    )

    h1(doc, "7. Keyword Strategy")
    h2(doc, "7.1 Rwanda (.rw) focus")
    bullets(
        doc,
        [
            "education in Rwanda, edtech Rwanda, education Kigali",
            "university software Rwanda, school management Rwanda",
            "MoMo school fees, MTN MoMo education, Vision 2050 education",
            "ICT education Rwanda, online learning Rwanda, StackEDU Rwanda",
            "Plus shared terms: school management system, SIS, AI powered edtech, academy, training",
        ],
    )
    h2(doc, "7.2 Africa (.africa) focus")
    bullets(
        doc,
        [
            "education in Africa, edtech Africa, African universities",
            "mobile money tuition Africa, M-Pesa school fees",
            "digital education Africa, AI education Africa, StackEDU Africa",
            "Same shared institutional / SIS / e-learning vocabulary",
        ],
    )

    h1(doc, "8. Local SEO (Rwanda)")
    bullets(
        doc,
        [
            "geo.region=RW, placename=Kigali, Rwanda",
            "GeoCoordinates latitude -1.9441 / longitude 30.0619",
            "LocalBusiness + PostalAddress (Kigali, RW)",
            "areaServed Country Rwanda on .rw; Place Africa on .africa",
            "Contact points with EN/FR/RW languages on .rw",
        ],
    )

    h1(doc, "9. Search Console & Verification")
    bullets(
        doc,
        [
            "stackedu.rw — HTML file: /google63477aeb3b4ade3c.html (middleware-excluded)",
            "stackedu.africa — HTML meta tag verification token in root metadata",
            "Crawler geo-bypass ensures Google fetches the correct property",
        ],
    )

    h1(doc, "10. Expert SEO Tester Checklist")
    h2(doc, "10.1 Automated / post-deploy checks")
    numbered(
        doc,
        [
            "curl https://www.stackedu.rw/robots.txt — Host + Sitemap point to .rw",
            "curl https://www.stackedu.africa/robots.txt — Host + Sitemap point to .africa",
            "curl https://www.stackedu.rw/sitemap.xml — all URLs on stackedu.rw",
            "curl https://www.stackedu.africa/sitemap.xml — all URLs on stackedu.africa",
            "curl https://www.stackedu.rw/llms.txt — Rwanda keywords + stackedu.rw canonical",
            "View source / — JSON-LD @graph present; title/description Rwanda-focused on .rw",
            "View source .africa / — Africa-focused copy + Africa verification meta",
            "Google Rich Results Test on home URL for both domains",
            "Search Console → URL Inspection → “Live URL” for /, /next, /sitemap.xml",
            "Confirm Googlebot (US) is NOT redirected away from .rw for verification/SEO assets",
        ],
    )
    h2(doc, "10.2 Ranking expectations (honest)")
    para(
        doc,
        "Technical SEO alone cannot guarantee #1 rankings overnight. This implementation "
        "removes crawl/indexation blockers, establishes topical and local entity signals, "
        "and positions StackEDU to compete for Rwanda and Africa edtech queries. Sustained "
        "rankings also require content velocity (blog), backlinks, citations, and Search "
        "Console performance monitoring.",
    )
    h2(doc, "10.3 Recommended next growth actions (outside this code drop)")
    bullets(
        doc,
        [
            "Submit both sitemaps in Google Search Console and Bing Webmaster Tools",
            "Publish Rwanda-focused blog posts targeting “school management system Rwanda”, “edtech Kigali”, etc.",
            "Earn .gov.rw / university / partner citations and backlinks",
            "Claim/consistent NAP citations for StackForgeAI Kigali",
            "Monitor GSC queries weekly; expand FAQ schema from real search questions",
            "Add OG image variants (1200×630) branded per market when design assets are ready",
        ],
    )

    h1(doc, "11. DRY Architecture Notes")
    para(
        doc,
        "All market SEO differences flow from getMarketSeo(market) and page copy helpers. "
        "Pages call buildPageMetadata() + JsonLd builders — no duplicated title/description "
        "strings across routes. Robots, sitemap, manifest, and llms.txt all reuse the same "
        "request-market helper.",
    )

    h1(doc, "12. Endpoints Reference")
    bullets(
        doc,
        [
            "https://stackedu.rw/ and https://stackedu.africa/",
            "https://stackedu.rw/robots.txt · https://stackedu.africa/robots.txt",
            "https://stackedu.rw/sitemap.xml · https://stackedu.africa/sitemap.xml",
            "https://stackedu.rw/llms.txt · https://stackedu.africa/llms.txt",
            "https://stackedu.rw/manifest.webmanifest · https://stackedu.africa/manifest.webmanifest",
            "https://stackedu.rw/google63477aeb3b4ade3c.html",
            "https://stackedu.rw/next · https://stackedu.africa/next",
        ],
    )

    h1(doc, "13. Conclusion")
    para(
        doc,
        "StackEDU now has a production-grade, dual-domain SEO foundation: complete metadata, "
        "structured data, sitemaps, robots, AI discoverability (llms.txt), local Rwanda signals, "
        "and Africa-wide alternate targeting. After deploy, verify endpoints, submit sitemaps, "
        "and continue content/backlink execution to convert technical readiness into rankings.",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("© StackForgeAI · StackEDU SEO Implementation Report · Confidential")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
